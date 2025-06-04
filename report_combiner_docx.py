from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_page_break_doc():
    """Creates a temporary document with a page break."""
    temp_doc = Document()
    paragraph = temp_doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return temp_doc

def combine_all_docx(filename_master, files_list):
    master = Document_compose(filename_master)
    composer = Composer(master)

    # Добавим разрыв страницы после мастер-документа
    composer.append(create_page_break_doc())

    for i, file_path in enumerate(files_list):
        print(file_path)
        doc_temp = Document_compose(file_path)
        composer.append(doc_temp)

        # Вставляем разрыв страницы после каждого документа, кроме последнего
        if i < len(files_list) - 1:
            composer.append(create_page_break_doc())

    composer.save("combined_file.docx")



files_to_merge = [
    "mount/src/glavpro/termins_and_short/dictionary_table.docx",
    "mount/src/glavpro/introduction/introduction.docx",
    "mount/src/glavpro/inventarization_description\punkt1\organization_sources_info.docx",
    "object_property\sanitary_zone\sanitary_zone.docx",
    r"inventarization_description\izav_info\razdel2.docx",
    r"calculations\tables\razdel3\razdel3.docx"
]
# filename_master="report.docx"
filename_master="content/Содержание_отчета.docx"
combine_all_docx(filename_master,files_to_merge)