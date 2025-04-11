import pandas as pd
from docx import Document

doc = Document()

data = pd.read_excel("company_info/новая_таблица.xlsx")

data.head()


from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Создаем новый документ
doc = Document()
style = doc.styles['Normal']

style.font.name = 'Times New Roman'  # Название шрифта
style.font.size = Pt(8)  # Размер шрифта (12 пунктов)
doc.add_heading('Отчёт об инвентаризации выбросов', 0)

# Change the page orientation to landscape
section = doc.sections[0]
section._sectPr.xpath('./w:pgSz')[0].set(qn('w:orient'), 'landscape')

# Adjust page width and height for landscape orientation
section.page_width = Inches(11)  # Landscape width (11 inches)
section.page_height = Inches(8.5)  # Landscape height (8.5 inches)
table = doc.add_table(rows=4, cols=18) # это число нормальное ставить
table.style = 'Table Grid'

# Объединяем ячейки для заголовков
table.cell(0, 0).text = "Наименование данных" # N цеха
table.cell(0, 1).text = "На момент разработки отчета инвентаризации"


for i in range(1, len(data)):
    table.add_row()
    table.cell(i, 0).text = data.loc[i - 1][0] # N цеха
    table.cell(i, 1).text = data.loc[i - 1][1]

doc.save('company_info/org_info.docx')




print(data.loc[0][1])