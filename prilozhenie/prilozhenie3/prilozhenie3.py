from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def generate_emissions_report(output_pdf):
    doc = Document()

    # Add a centered title
    title = doc.add_paragraph("Результаты инструментального определения показателей выбросов")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Add school information
    school_info = [
        "Муниципальное бюджетное общеобразовательное учреждение",
        "«Мирошкинская средняя общеобразовательная школа»",
        "Первомайского района Оренбургской области",
        "(МБОУ «Мирошкинская СОШ»)"
    ]

    for line in school_info:
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(12)

    # Add the main text
    doc.add_paragraph("\n")  # Add spacing
    text = doc.add_paragraph(
        "Аналитические лаборатории не привлекались. Инвентаризация источников загрязнения, "
        "количественная и качественная характеристика выбросов осуществлялись при помощи "
        "расчётного (балансового) метода."
    )
    text.runs[0].font.size = Pt(12)

    # Save as DOCX first
    temp_docx = "emissions_report.docx"
    doc.save(temp_docx)

    # Convert to PDF
    convert(temp_docx, output_pdf)
    print(f"PDF generated successfully: {output_pdf}")

# Example usage
generate_emissions_report("prilozhenie/prilozhenie3/prilozhenie3.pdf")