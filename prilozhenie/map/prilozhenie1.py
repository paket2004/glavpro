from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_school_map_document(image_path, output_pdf_path):
    # Создаем новый документ Word
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Добавляем заголовок "КАРТА-СХЕМА"
    heading = doc.add_paragraph('КАРТА-СХЕМА')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.bold = True
    heading.runs[0].font.size = Pt(14)
    
    # Добавляем подзаголовок
    doc.add_paragraph('территории предприятия МБОУ "Мирошкинская СОШ"' \
    '461987, Оренбургская область, Первомайский район, с. Мирошкино, ул. Центральная, 27')
    
    # Добавляем пустую строку для отступа
    doc.add_paragraph()
    
    # Вставляем изображение
    doc.add_picture(image_path, width=Inches(6))  # Ширина 6 дюймов (можно регулировать)
    
    # Сохраняем документ Word
    docx_path = os.path.splitext(output_pdf_path)[0] + '.docx'
    doc.save(docx_path)
    
    # Конвертируем в PDF (требуется установленный Word или библиотека comtypes)
    try:
        from docx2pdf import convert
        convert(docx_path, output_pdf_path)
        print(f"Документ успешно сохранен как PDF: {output_pdf_path}")
    except ImportError:
        print(f"Сохранен документ Word: {docx_path}. Для конвертации в PDF установите docx2pdf (pip install docx2pdf)")

# Использование
image_path = r"prilozhenie/map/annotated_map/map_schema.jpg"
output_pdf = r"prilozhenie/map/prilozhenie1.pdf"
create_school_map_document(image_path, output_pdf)