from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(6)
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('''
    Таблица № 1.1. 
    Режимы работы ИЗАВ и их временные характеристики 
    при нестационарности выбросов''')
run.font.color.rgb = RGBColor(0,0,0)
run.font.size = Pt(14)
table = doc.add_table(rows=4, cols=6)
table.style = 'Table Grid'

headers = [
    "№ ИЗАВ", "", "", "", "Тип ИЗАВ", "N (код) режима ИЗАВ (в зависимости от времени работы ИВ, одинаковый для одновременно работающих ИЗАВ)"
]
# Объединяем ячейки для заголовков
table.cell(0, 0).merge(table.cell(1, 0))
table.cell(0, 1).merge(table.cell(0, 4))
table.cell(1, 1).text = "Номер ИВ"
table.cell(1, 2).text = "Наименование ИВ"
table.cell(1, 3).text = "Описание режима работы ИВ"
table.cell(1, 4).text = "Время работы ИВ на конкретном режиме за период времени"
table.cell(0, 5).merge(table.cell(1, 5))
for i, header in enumerate(headers):
    table.cell(0, i).text = header

table.cell(2,0).merge(table.cell(2,5))
table.cell(2,0).text = "На предприятии один (основной) режим работы, нестационарности не выявлено"
doc.save('prilozhenie/tables/1_1.docx')
