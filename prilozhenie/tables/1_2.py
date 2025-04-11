from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(6)
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('''
    Таблица № 1.2.
    Характеристика одновременности работы оборудования
    при нестационарных выбросах''')
run.font.color.rgb = RGBColor(0,0,0)
run.font.size = Pt(14)
table = doc.add_table(rows=4, cols=8)
table.style = 'Table Grid'
headers = [
    "Наименование цеха", "", "", "", "", "Источники выделения (выброса)", "Коэффициент К0", "Номер ИЗАВ"
]
table.cell(0, 0).merge(table.cell(2, 0))

table.cell(0, 1).merge(table.cell(0, 5))
table.cell(1, 1).text = "№№"
table.cell(1, 2).text = "Наименование"
table.cell(1, 3).text = "Режим ИВ"
table.cell(1, 4).merge(table.cell(1, 5))
table.cell(1, 4).text = "Количество"
table.cell(2, 4).text = "Всего" 
table.cell(2, 5).text = "В т.ч. одновременно работающих"
for i, header in enumerate(headers):
    table.cell(0, i).text = header

table.cell(0,6).merge(table.cell(2,6))
table.cell(0,7).merge(table.cell(2,7))
doc.save('prilozhenie/tables/1_2.docx')
