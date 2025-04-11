from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(6)
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('''
    Таблица № 1.3.
    Учет нестационарности выбросов''')
run.font.color.rgb = RGBColor(0,0,0)
run.font.size = Pt(14)
table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'

headers = [
    "№", "№ ИЗАВ", "Источник выделения", "Характеристики технологических стадий", "Значения характеристик технологических стадий"
]
# Объединяем ячейки для заголовков
table.cell(0, 0).merge(table.cell(2, 0))  # N ИЗАВ
table.cell(0, 1).merge(table.cell(3, 1))  # Количество ЗВ, отходящих от ИВ
table.cell(0, 2).merge(table.cell(3, 2))
table.cell(0, 3).merge(table.cell(1, 3))
table.cell(0, 4).merge(table.cell(2, 4))
# table.cell(0, 5).merge(table.cell(2, 5))
table.cell(2,3).text = "Название характеристики"
# table.cell(1, 3).merge(table.cell(2, 3))
for i, header in enumerate(headers):
    table.cell(0, i).text = header

    # # Добавляем пустые строки для данных
    # for row in range(2, 3):  # Добавляем одну пустую строку для данных
    #     for col in range(18):
    #         table.cell(row, col).text = ""  # Оставляем ячейки пустыми
    # Сохраняем документ
# Стадии технологического процесса отсутствуют или не описаны!
table.cell(3,0).merge(table.cell(3,4))
table.cell(3,0).text = "Стадии технологического процесса отсутствуют или не описаны!"
doc.save('prilozhenie/tables/1_3.docx')
