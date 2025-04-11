from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document()

par = doc.add_paragraph()
par.add_run('Таблицы учёта характеристик нестационарности выбросов')
par.alignment = WD_ALIGN_PARAGRAPH.CENTER

par = doc.add_paragraph()
run = par.add_run('''
    Муниципальное бюджетное общеобразовательное учреждение
    «Мирошкинская средняя общеобразовательная школа»
    Первомайского района Оренбургской области
    (МБОУ «Мирошкинская СОШ»)''')
run.bold = True
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save("prilozhenie/tables/intro_pr4.docx")