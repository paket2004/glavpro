from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def generate_emissions_report(output_pdf):
    doc = Document()

    # Add a centered title
    title = doc.add_paragraph('''
    Копия аттестата аккредитации привлекаемой аналитической
    лаборатории с приложением области аккредитации, копии
    материалов, использованных в ходе инвентаризации выбросов и
    составления отчета''')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("\n")  # Add spacing
    text = doc.add_paragraph(
        "Аналитические лаборатории не привлекались. Инвентаризация источников загрязнения, "
        "количественная и качественная характеристика выбросов осуществлялись при помощи "
        "расчётного (балансового) метода."
    )
    text.runs[0].font.size = Pt(12)

    # Save as DOCX first
    temp_docx = "emissions_report.docx"
    doc.save(temp_docx)

    # Convert to PDF
    convert(temp_docx, output_pdf)
    print(f"PDF generated successfully: {output_pdf}")

# Example usage
generate_emissions_report("prilozhenie/prilozhenie5/prilozhenie5.pdf")