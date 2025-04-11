import os
import sys
sys.path.append(os.getcwd())

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from gross_emissions.calculate_gross_emissions import calculate_gross_emissions

# calculate_gross_emissions
def generate_moving_desription(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk):
    _, COgik_warm, CO_mik_warm, _, Gno2_warm, _, Gno_warm, _, Mno2_warm, Mno_warm, _, _, _, SO_gik_warm, _, SO_mik_warm, petrol_gik_warm, _, petrol_mik_warm, _ = calculate_gross_emissions(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    doc = Document()
    
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = heading.add_run('РАСЧЁТ ВАЛОВЫХ ВЫБРОСОВ')
    run.bold = True
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(14)

    doc.add_paragraph('''
    Площадка 01
    Стационарный источник загрязнения 6003, режим ИЗАВ: 1
    Передвижной источник загрязнения: Движение и работа транспорта по территории (автобус)''')
    
    doc.add_paragraph('''
    РАСЧЕТ ВЫБРОСОВ ЗАГРЯЗНЯЮЩИХ ВЕЩЕСТВ 
    ПРИ ДВИЖЕНИИ АВТОМОБИЛЕЙ ПО ТЕРРИТОРИИ''')

    doc.add_paragraph('''
    1. Расчет выбросов от различных групп автомобилей ведется по "Методике
    проведения инвентаризация выбросов загрязняющих веществ в атмосферу
    для автотранспортных предприятий". М,1998.п2., с учетом дополнений 1999 г.''')
    
    doc.add_paragraph('''
    2. Расчет выбросов от дорожных машин ведется по "Методике
    проведения инвентаризация выбросов загрязняющих веществ в атмосферу для баз
    дорожной техники". М,1998.п2.''')
    
    doc.add_paragraph('''
                      
        Выброс загрязняющих веществ одним автомобилем данной группы
    в день при движении по территории предприятия
    рассчитывается c использованием формулы (2.11) из [1]''')
    
    par = doc.add_paragraph()
    par.add_run("""
                M1iк = mLiк · L1, г (1)
                """).bold = True
    
    doc.add_paragraph("""
        где mLiк - пробеговый выброс вещества автомобилем при движении
        по территории предприятия, г/км
        L1 - пробег автомобиля по территории предприятия, км/день""")
    
    doc.add_paragraph('''
        Максимальный выброс от 1 автомобиля данной группы
    в течении периода времени Tr рассчитывается c использованием формулы (2.13) из [1]:''')
    
    par = doc.add_paragraph()
    par.add_run('''M2iк = mLiк · L2 , г (2)''').bold = True
    
    doc.add_paragraph('''
        где L2 - максимальный пробег автомобиля за Tr мин, км
            Tr - период времени в минутах, характеризующийся максимальной
        интенсивностью движения автотранспорта по
        территории предприятия''')
     

    doc.add_paragraph('''
        Валовый выброс вещества автомобилями данной
    группы рассчитывается раздельно для каждого периода по формуле (2.11) из [1]:''')

    par = doc.add_paragraph()
    par.add_run("""Miк = M1iк · Nкв · Dp · 10-6 , т / год (3)""").bold = True
    doc.add_paragraph('''
        где Nкв - среднее количество автомобилей данной группы,
        двигающихся по территории предприятия в сутки
        Dp - количество рабочих дней в расчетном периоде (теплый, переходный, холодный)''')
    
    doc.add_paragraph('''
        Для определения общего валового выброса валовые выбросы одноименных
        веществ от разных групп автомобилей и разных расчетных периодов года
        суммируются''')
    
    doc.add_paragraph('''
        Максимально разовый выброс от автомобилей
    данной группы рассчитывается по формуле:''')
    par = doc.add_paragraph()
    par.add_run('''Giк = M2iк · N'к / Tr / 60 , г / c (4)''').bold = True

    doc.add_paragraph('''
    где N'к - наибольшее количество машин данной группы, двигающихся
    в течение периода времени Tr минут''')

    doc.add_paragraph('''
    Из полученных значений G для разных групп автомобилей и расчетных
    периодов выбирается максимальное.
    Если одновременно двигаются автомобили разных групп,
    то их разовые выбросы суммируются''')

    doc.add_paragraph('''
    ______________________________________________________________________
    Коэффициент трансформации окислов азота в NO2, kno2 = 0.8
    Коэффициент трансформации окислов азота в NO, kno = 0.13
                      
    ______________________________________________________________________
    Расчетный период: Переходный период (t> = -5 и t< = 5)
    ______________________________________________________________________
    Температура воздуха за расчетный период, град. С, t = 5
    Период максимальной интенсивности движения техники по территории п/п, мин, Tr = 20
    ______________________________________________________________________
    Тип машины: Автобусы карбюраторные особо малые габаритной длиной до 5.5 м (СНГ)
    ______________________________________________________________________
    Тип топлива: Бензин А-76, АИ-92
    Экологический контроль не проводится''')    
        
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    table.cell(0,0).text = "Dp, сут"
    table.cell(0,1).text = "Nk, шт."
    table.cell(0,2).text = "Nkв, шт."
    table.cell(0,3).text = "N'k, шт."
    table.cell(0,4).text = "L1, км"
    table.cell(0,5).text = "L2, км"

    table.cell(1,0).text = f"{Dp}"
    table.cell(1,1).text = f"{Nk}"
    table.cell(1,2).text = f"{Nkv}"
    table.cell(1,3).text = f"{Nkk}"
    table.cell(1,4).text = f"{L1_warm}"
    table.cell(1,5).text = f"{L2_warm}"
    
    doc.add_paragraph()
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0337 Углерода оксид''')
    run.bold = True
    run.underline = True
    # вот эти данные по хорошему надо как-то хранить и откуда-то подтягивать, а не каждый раз вручную смотреть в методичке
    par = doc.add_paragraph()
    run = par.add_run(f'''
    mLiк = 25.65
    mxxiк = 4.5
    M1iк = mLiк · {L1_warm} = 25.65 · 0.1 = 2.565
    Miк = aв · M1iк · Nk · Dp · 10-6 = 1 · 2.565 · 1 · 181 · 10-6 = {CO_mik_warm}
    M2iк = mLiк · L2 = 25.65 · {L2_warm} = 2.565
    Giк = M2iк · N'к / Tr / 60 = 2.565 · {Nk} / {Tr} / 60 = {COgik_warm}''')
   
    
    
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 2704 Бензин (нефтяной, малосернистый) /в пересчете на углерод/''')
    run.bold = True
    run.underline = True
    
    par = doc.add_paragraph()
    par.add_run(f'''
    mLiк = 3.15
    mxxiк = 0.4
    M1iк = mLiк · {L1_warm} = 3.15 · 0.1 = 0.315
    Miк = aв · M1iк · {Nk} · {Dp} · 10-6 = 1 · 0.315 · 1 · 181 · 10-6 = {petrol_mik_warm}
    M2iк = mLiк · L2 = 3.15 · {L2_warm} = 0.315
    Giк = M2iк · N'к / Tr / 60 = 0.315 · {Nkk} / {Tr} / 60 = {petrol_gik_warm}''').bold = True

    doc.add_paragraph('''РАСЧЕТ выбросов оксидов азота:''')
    par = doc.add_paragraph()
    par.add_run(f'''
    mLiк = 0.6
    mxxiк = 0.05
    M1iк = mLiк · L1 = 0.6 · {L1_warm} = 0.06
    Miк = aв · M1iк · {Nk} · {Dp} · 10-6 = 1 · 0.06 · 1 · 181 · 10-6 = 0.00001086
    M2iк = mLiк · L2 = 0.6 · {L2_warm} = 0.06
    Giк = M2iк · Nk / Tr / 60 = 0.06 · {Nk} / {Tr} / 60 = 0.00005''').bold = True

    doc.add_paragraph('''С учетом трансформации оксидов азота получаем:''')
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0301 Азота диоксид''')
    run.bold = True
    run.underline = True

    doc.add_paragraph(f'''Валовый выброс, т/год, Mno2 = kno2 · Miк = 0.8 · 0.00001086 = {Mno2_warm}''')
    doc.add_paragraph(f'''Максимальный разовый выброс,г/с, Gno2 = kno2 · Giк = 0.8 · 0.00005 = {Gno2_warm}''')

    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0304 Азот (II) оксид''')
    run.bold = True
    run.underline = True
    doc.add_paragraph(f'''Валовый выброс, т/год, Mno = kno · Miк = 0.13 · 0.00001086 ={Mno_warm}''')
    doc.add_paragraph(f'''Максимальный разовый выброс,г/с, Gno = kno · Giк = 0.13 · 0.00005 = {Gno_warm}''')

    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0330 Сера диоксид''')
    run.bold = True
    run.underline = True
    par = doc.add_paragraph()
    par.add_run(f'''
    mLiк = 0.099
    mxxiк = 0.012
    M1iк = mLiк · L1 = 0.099 · {L1_warm} = 0.0099
    Miк = aв · M1iк · Nk · Dp · 10-6 = 1 · 0.0099 · {Nk} · {Dp} · 10-6 = {SO_mik_warm}
    M2iк = mLiк · L2 = 0.099 · 0.1 = 0.0099
    Giк = M2iк · N'к / Tr / 60 = 0.0099 · {Nkk} / {Tr} / 60 = {SO_gik_warm}''').bold = True
    
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'  
    table.cell(0,0).text = "Код ЗВ"
    table.cell(0,1).text = "Наименование ЗВ"
    table.cell(0,2).text = "mLiк, г/км"
    table.cell(0,3).text = "G, г/с"
    table.cell(0,4).text = "M, т/г"

    table.cell(1,0).text = "0337"
    table.cell(1,1).text = "Углерода оксид"
    table.cell(1,2).text = "25.65"
    table.cell(1,3).text = f"{COgik_warm}"
    table.cell(1,4).text = f"{CO_mik_warm}"

    table.cell(2,0).text = "2704"
    table.cell(2,1).text = "Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(2,2).text = "3.15"
    table.cell(2,3).text = f"{petrol_gik_warm}"
    table.cell(2,4).text = f"{petrol_mik_warm}"

    table.cell(3,0).text = "0301"
    table.cell(3,1).text = "Азота диоксид"
    table.cell(3,2).text = "0.6"
    table.cell(3,3).text = f"{Gno2_warm}"
    table.cell(3,4).text = f"{Mno2_warm}"
    
    table.cell(4,0).text = "0304"
    table.cell(4,1).text = "Азот (II) оксид"
    table.cell(4,2).text = "0.6"
    table.cell(4,3).text = f"{Gno_warm}"
    table.cell(4,4).text = f"{Mno_warm}"

    table.cell(5,0).text = "0330"
    table.cell(5,1).text = "Сера диоксид"
    table.cell(5,2).text = "0.099"
    table.cell(5,3).text = f"{SO_gik_warm}"
    table.cell(5,4).text = f"{SO_mik_warm}"


    doc.add_paragraph('''
    ИТОГО выбросы по периоду: Переходный период (t> = -5 и t< = 5)
    Температура воздуха за расчетный период, град. С, t = 5''')

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    table.cell(0,0).merge(table.cell(0,3))
    table.cell(0,0).text = "ВСЕГО по периоду: Переходный период (t>=-5 и t<=5)"
    
    table.cell(1,0).text = "Код"
    table.cell(1,1).text = "Наименование ЗВ"
    table.cell(1,2).text = "Выброс г/с"
    table.cell(1,3).text = "Выброс т/год"

    table.cell(2,0).text = "0337"
    table.cell(2,1).text = "Углерода оксид"
    table.cell(2,2).text = f"{COgik_warm}"
    table.cell(2,3).text = f"{CO_mik_warm}"

    table.cell(3,0).text = "2704"
    table.cell(3,1).text = "Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(3,2).text = f"{petrol_gik_warm}"
    table.cell(3,3).text = f"{petrol_mik_warm}"

    table.cell(4,0).text = "0301"
    table.cell(4,1).text = "Азота диоксид"
    table.cell(4,2).text = f"{Gno2_warm}"
    table.cell(4,3).text = f"{Mno2_warm}"

    table.cell(5,0).text = "0330"
    table.cell(5,1).text = "Сера диоксид"
    table.cell(5,2).text = f"{SO_gik_warm}"
    table.cell(5,3).text = f"{SO_mik_warm}"

    table.cell(6,0).text = "0304"
    table.cell(6,1).text = "Азот (II) оксид"
    table.cell(6,2).text = f"{Gno_warm}"
    table.cell(6,3).text = f"{Gno2_warm}"



    doc.add_paragraph('''ИТОГО ВЫБРОСЫ''')

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.cell(0,0).text = "Код"
    table.cell(0,1).text = "Наименование ЗВ"
    table.cell(0,2).text = "Выброс г/с"
    table.cell(0,3).text = "Выброс т/год"


    table.cell(1,0).text = "0301"
    table.cell(1,1).text = "Азота диоксид"
    table.cell(1,2).text = f"{Gno2_warm}"
    table.cell(1,3).text = f"{Mno2_warm}"

    table.cell(2,0).text = "0304"
    table.cell(2,1).text = "Азот (II) оксид"
    table.cell(2,2).text = f"{Gno_warm}"
    table.cell(2,3).text = f"{Mno_warm}"

    table.cell(3,0).text = "0330"
    table.cell(3,1).text = "Сера диоксид"
    table.cell(3,2).text = f"{SO_gik_warm}"
    table.cell(3,3).text = f"{SO_mik_warm}"

    table.cell(4,0).text = "0337"
    table.cell(4,1).text = "Углерода оксид"
    table.cell(4,2).text = f"{COgik_warm}"
    table.cell(4,3).text = f"{CO_mik_warm}"

    table.cell(5,0).text = "2704"
    table.cell(5,1).text = "Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(5,2).text = f"{petrol_gik_warm}"
    table.cell(5,3).text = f"{petrol_mik_warm}"
    doc.add_paragraph()
    doc.add_paragraph('''Максимально-разовые выбросы достигнуты в переходный период''')
    doc.save('prilozhenie/prilozhenie2/moving.docx')
    print("sth")
# generate_moving_desription()
