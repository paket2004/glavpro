from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from gross_emissions.calculate_gross_emissions import calculate_gross_emissions
def generate_station_desription(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk):
    COgik_warmup, _, _, CO_mik_warmup, _, Gno2_warm_up, _, Gno_warm_up, _, _, Mno_warm_up, Mno2_warm_up, SOgik, _, SO_mik, _, _, petrol_gik_warmup, _, petrol_mik_warmup = calculate_gross_emissions(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    doc = Document()
    
    # doc.add_heading('РАСЧЁТ ВАЛОВЫХ ВЫБРОСОВ')
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = heading.add_run('РАСЧЁТ ВАЛОВЫХ ВЫБРОСОВ')
    run.bold = True
    run.font.color.rgb = RGBColor(0,0,0)
    run.bold = True
    run.font.size = Pt(10)  # Размер 10 для заголовка


    
    doc.add_paragraph('''
    Площадка 01
    Стационарный источник загрязнения 6003, режим ИЗАВ: 1
    Источник выделения: 001, Открытая стоянка''')
    
    doc.add_paragraph('''
    РАСЧЕТ ВЫБРОСОВ ЗАГРЯЗНЯЮЩИХ ВЕЩЕСТВ
    ОТ СТОЯНОК АВТОМОБИЛЕЙ''')

    doc.add_paragraph('''
    1. Расчет выбросов от различных групп автомобилей ведется по "Методике
    проведения инвентаризация выбросов загрязняющих веществ в атмосферу для автотранспортных предприятий". М,1998.п2., с учетом дополнений 1999 г.''')
    
    doc.add_paragraph('''
    2. Расчет выбросов от дорожных машин ведется по "Методике
    проведения инвентаризация выбросов загрязняющих веществ в атмосферу для баз дорожной техники". М,1998.п2.''')
    
    doc.add_paragraph('''
    Выброс загрязняющих веществ одним автомобилем данной группы в день 
    при выезде с территории или помещения стоянки (M1ik) и возврате (M2ik)
    расчитывается по формулам (2.1), (2.2), из [1]: (расчетная схема 1)''')
    

    par = doc.add_paragraph()
    par.add_run("""
        M1ik = mnpik * tnp + mLik * L1 + mxxik * txx1, г           (1)
        M2ik = mLik * L2 + mxxik * txx2, г                (2)""").bold = True
    
    doc.add_paragraph('''
        Где mnpik - удельный выброс вещества при прогреве двигателя автомобиля, г/мин.
        mLik - пробеговый выброс вещества автомобилем, г/км                  
        mxxik -  удельный выброс вещества при работе двигателя на холостом ходу, г/мин 
        tnp - время прогрева двигателя, мин                  
        txx1, txx2 - время работы двигателя на холостом ходу при выезде и возврате. txx2 = txx1 = 1 мин.
        L1, L2 - пробег автомобиля по территории стоянки, км
    ''')

    doc.add_paragraph("""
    Валовый выброс вещества автомобилями данной группы рассчитывается
    раздельно для каждого периода по формуле (2.7) из [1]:""")
    
    par = doc.add_paragraph()
    par.add_run('''     Miк = aв · (M1iк + M2iк) · Nk · Dp · 10-6, т / год          (3)''').bold = True
    doc.add_paragraph('''
        где aв - коэффициент выпуска (выезда), aв = Nкв/Nk
        Nкв - среднее количество автомобилей данной группы, выходящих со стоянки в сутки                  
        Nk - общее количество автомобилей данной группы на территории или в помещении стоянки
        Dp - количество рабочих дней в расчетном периоде (холодном, теплом, переходном)''')
    
    doc.add_paragraph('''
        Для определения общего валового выброса, валовые выбросы одноименных
    веществ по периодам года суммируются''')
    doc.add_paragraph('''
        Максимально разовый выброс вещества рассчитывается для каждого
    периода по формуле:
''')
    
    par = doc.add_paragraph()
    par.add_run('''     Giк = MAX(M1iк,M2iк) · N'к / Tr / 60, г / c (4)''').bold = True
    doc.add_paragraph('''
    где MAX(M1iк,M2iк) - максимум из выбросов вещества при выезде и въезде
    автомобиля данной группы, г
    Tr - период времени в минутах, характеризующийся максимальной
    интенсивностью выезда (въезда) автомобилей на стоянку
    N'к - наибольшее количество автомобилей данной группы, выезжающих
    со стоянки (въезжающих на стоянку) в течение периода времени Tr               
    ''')

    doc.add_paragraph('''
    Из полученных значений G для разных групп автомобилей и расчетных
    периодов выбирается максимальное.''')
    doc.add_paragraph('''
    Если в течение периода времени Tr выезжают (въезжают) автомобили разных
    групп, то их разовые выбросы суммируются.
''')
    
    doc.add_paragraph('''
    ______________________________________________________________________
    Коэффициент трансформации окислов азота в NO2, kno2 = 0.8
    Коэффициент трансформации окислов азота в NO, kno = 0.13
                                   
    Стоянка: Обособленная, имеющая непосредственный выезд на дорогу общего пользования (расчетная схема 1)              
    Условия хранения: Открытая или закрытая неотапливаемая стоянка без средств подогрева           
    ______________________________________________________________________''')

    doc.add_paragraph('''
    Расчетный период: Переходный период (t> = -5 и t< = 5)               
    ______________________________________________________________________               
    Температура воздуха за расчетный период, град. С, t = 5             
    Период максимальной интенсивности выезда техники со стоянки, мин, Tr = 20           
    ______________________________________________________________________        
    Тип машины: Автобусы карбюраторные особо малые габаритной длиной до 5.5 м (СНГ)  
    ______________________________________________________________________''')

    doc.add_paragraph('''
    Тип топлива: Бензин А-76, АИ-92
    Экологический контроль не проводится''')
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    table.cell(0,0).text = "Dp, сут"
    table.cell(0,1).text = "Nk, шт."
    table.cell(0,2).text = "Nkв, шт."
    table.cell(0,3).text = "N'k, шт."
    table.cell(0,4).text = "L1, км"
    table.cell(0,5).text = "L2, км"

    table.cell(1,0).text = f"{Dp}"
    table.cell(1,1).text = f"{Nk}"
    table.cell(1,2).text = f"{Nkv}"
    table.cell(1,3).text = f"{Nkk}"
    table.cell(1,4).text = f"{L1_warmup}"
    table.cell(1,5).text = f"{L2_warmup}"

    doc.add_paragraph()
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0337 Углерода оксид''')
    run.bold = True
    run.underline = True
    par = doc.add_paragraph()
    par.add_run(f'''
    mпрiк = 8.19
    mLiк = 25.65
    mxxiк = 4.5
    M1iк = mпрiк · tпр + mLiк · L1 + mxxiк · txx1 = 8.19 · {tnp} + 25.65 · {L1_warmup} + 4.5 · {txx1} = 37.5
    M2iк = mLiк · L2 + mxxiк · txx2 = 25.65 · {L2_warmup} + 4.5 · {txx2} = 4.76
    Miк = aв · (M1iк + M2iк) · Nk · Dp · 10-6 = {ab} · (37.5 + 4.76) · {Nk} · {Dp} · 10-6 = {CO_mik_warmup}                                   
    Giк = (mпрiк · tпр + mLiк · L1 + mxxiк · txx1) · N'к / Tr / 60 = (8.19 · {tnp} + 25.65 · {L1_warmup} + 4.5 · {txx1}) · {Nkk} / {Tr} / 60 = {COgik_warmup}
    ''').bold = True
    
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 2704 Бензин (нефтяной, малосернистый) /в пересчете на углерод/''')
    run.bold = True
    run.underline = True
    par = doc.add_paragraph()
    par.add_run(f'''
    mпрiк = 0.9
    mLiк = 3.15
    mxxiк = 0.4
    M1iк = mпрiк · tпр + mLiк · L1 + mxxiк · txx1 = 0.9 · {tnp} + 3.15 · {L1_warmup} + 0.4 · {txx1} = 4.03 
    M2iк = mLiк · L2 + mxxiк · txx2 = 3.15 · {L2_warmup} + 0.4 · {txx2} = 0.4315 
    Miк = aв · (M1iк + M2iк) · Nk · Dp · 10-6 = {ab} · (4.03 + 0.4315) · {Nk} · {Dp} · 10-6 = {petrol_mik_warmup}
    Giк = (mпрiк · tпр + mLiк · L1 + mxxiк · txx1) · N'к / Tr / 60 = (0.9 · {tnp} + 3.15 · {L1_warmup} + 0.4 · {txx1}) · {Nkk} / {Tr} / 60 = {petrol_gik_warmup}
    ''').bold = True
    
    doc.add_paragraph('''РАСЧЕТ выбросов оксидов азота:''')

    par = doc.add_paragraph()
    par.add_run(f'''
    mпрiк = 0.07
    mLiк = 0.6
    mxxiк = 0.05
    M1iк = mпрiк · tпр + mLiк · L1 + mxxiк · txx1 = 0.07 · {tnp} + 0.6 · {L1_warmup} + 0.05 · {txx1} = 0.336
    M2iк = mLiк · L2 + mxxiк · txx2 = 0.6 · {L2_warmup} + 0.05 · {txx2} = 0.056
    Miк = aв · (M1iк + M2iк) · Nk · Dp · 10-6 = {ab} · (0.336 + 0.056) · {Nk} · {Dp} · 10-6 = 0.000071
    Giк = (mпрiк · tпр + mLiк · L1 + mxxiк · txx1) · N'к / Tr / 60 = (0.07 · {tnp} + 0.6 · {L1_warmup} + 0.05 · {txx1}) · {Nkk} / {Tr} / 60 = 0.00028''').bold = True
   
    doc.add_paragraph('''С учетом трансформации оксидов азота получаем:''')

    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0301 Азота диоксид''')
    run.bold = True
    run.underline = True
    doc.add_paragraph('''
    Валовый выброс, т/год, Mno2 = kno2 · Miк = 0.8 · 0.000071 = 0.0000568
    Максимальный разовый выброс,г/с, Gno2 = kno2 · Giк = 0.8 · 0.00028 = 0.000224''')
    
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0304 Азот (II) оксид''')
    run.bold = True
    run.underline = True
    doc.add_paragraph('''
    Валовый выброс, т/год, Mno = kno · Miк = 0.13 · 0.000071 = 0.00000923
    Максимальный разовый выброс,г/с, Gno = kno · Giк = 0.13 · 0.00028 = 0.0000364''')
    
    doc.add_paragraph('''''')
    par = doc.add_paragraph()
    run = par.add_run('''Примесь: 0330 Сера диоксид''')
    run.bold = True
    run.underline = True
    par = doc.add_paragraph()
    run = par.add_run(f'''
    mпрiк = 0.0144
    mLiк = 0.099
    M1iк = mпрiк · tпр + mLiк · L1 + mxxiк · txx1 = 0.0144 · {tnp} + 0.099 · {L1_warmup} + 0.012 · {txx1} = 0.0706
    M2iк = mLiк · L2 + mxxiк · txx2 = 0.099 · {L2_warmup} + 0.012 · {txx2} = 0.013
    Miк = aв · (M1iк + M2iк) · Nk · Dp · 10-6 = {ab} · (0.0706 + 0.013) · {Nk} · {Dp} · 10-6 = 0.00001513
    Giк = (mпрiк · tпр + mLiк · L1 + mxxiк · txx1) · N'к / Tr / 60 = (0.0144 · {tnp} + 0.099 · {L1_warmup} + 0.012 · {txx1}) · {Nkk} / {Tr} / 60 = 0.0000588''')
    run.bold = True
    
    table = doc.add_table(rows=6, cols=9)
    table.style = 'Table Grid'

    table.cell(0,0).text = "Код ЗВ"
    table.cell(0,1).text = "Наименование ЗВ"
    table.cell(0,2).text = "tпр,мин"
    table.cell(0,3).text = "mпрiк, г/мин"
    table.cell(0,4).text = "txx1, мин"
    table.cell(0,5).text = "mxxiк, г/мин"
    table.cell(0,6).text = "mLiк, г/км"
    table.cell(0,7).text = "G, г/с"
    table.cell(0,8).text = "M, т/г"


    table.cell(1,0).text = "0337"
    table.cell(1,1).text = "Углерода оксид"
    table.cell(1,2).text = "4"
    table.cell(1,3).text = "8.19"
    table.cell(1,4).text = "1"
    table.cell(1,5).text = "4.5"
    table.cell(1,6).text = "25.65"
    table.cell(1,7).text = f"{COgik_warmup}"
    table.cell(1,8).text = f"{CO_mik_warmup}"

    table.cell(2,0).text = "2704"
    table.cell(2,1).text = "Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(2,2).text = "4"
    table.cell(2,3).text = "0.9"
    table.cell(2,4).text = "1"
    table.cell(2,5).text = "0.4"
    table.cell(2,6).text = "3.15"
    table.cell(2,7).text = f"{petrol_gik_warmup}"
    table.cell(2,8).text = f"{petrol_mik_warmup}"

    table.cell(3,0).text = "0301"
    table.cell(3,1).text = "Азота диоксид"
    table.cell(3,2).text = "4"
    table.cell(3,3).text = "0.07"
    table.cell(3,4).text = "1"
    table.cell(3,5).text = "0.05"
    table.cell(3,6).text = "0.6"
    table.cell(3,7).text = f"{Gno2_warm_up}"
    table.cell(3,6).text = f"{Mno2_warm_up}"
    
    table.cell(4,0).text = "0304"
    table.cell(4,1).text = "Азот (II) оксид"
    table.cell(4,2).text = "4"
    table.cell(4,3).text = "0.07"
    table.cell(4,4).text = "1"
    table.cell(4,5).text = "0.05"
    table.cell(4,6).text = "0.6"
    table.cell(4,7).text = f"{Gno_warm_up}"
    table.cell(4,8).text = f"{Mno_warm_up}"

    table.cell(5,0).text = "0330"
    table.cell(5,1).text = "Сера диоксид"
    table.cell(5,2).text = "4"
    table.cell(5,3).text = "0.014"
    table.cell(5,4).text = "1"
    table.cell(5,5).text = "0.012"
    table.cell(5,6).text = "0.099"
    table.cell(5,7).text = f"{SOgik}"
    table.cell(5,8).text = f"{SO_mik}"


    doc.add_paragraph('''
    ИТОГО выбросы по периоду: Переходный период (t> = -5 и t< = 5)
    Температура воздуха за расчетный период, град. С, t = 5''')

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    table.cell(0,0).merge(table.cell(0,3))
    table.cell(0,0).text = "ВСЕГО по периоду: Переходный период (t>=-5 и t<=5)"
    
    table.cell(1,0).text = "Код"
    table.cell(1,1).text = "Наименование ЗВ"
    table.cell(1,2).text = "Выброс г/с"
    table.cell(1,3).text = "Выброс т/год"

    table.cell(2,0).text = "0337"
    table.cell(2,1).text = "Углерода оксид"
    table.cell(2,2).text = f"{COgik_warmup}"
    table.cell(2,3).text = f"{CO_mik_warmup}"

    table.cell(3,0).text = "2704"
    table.cell(3,1).text = "Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(3,2).text = f"{petrol_gik_warmup}"
    table.cell(3,3).text = f"{petrol_mik_warmup}"

    table.cell(4,0).text = "0301"
    table.cell(4,1).text = "Азота диоксид"
    table.cell(4,2).text = f"{Gno2_warm_up}"
    table.cell(4,3).text = f"{Mno2_warm_up}"

    table.cell(5,0).text = "0330"
    table.cell(5,1).text = "Сера диоксид"
    table.cell(5,2).text = f"{SOgik}"
    table.cell(5,3).text = f"{SO_mik}"

    table.cell(6,0).text = "0304"
    table.cell(6,1).text = "Азот (II) оксид"
    table.cell(6,2).text = f"{Gno_warm_up}"
    table.cell(6,3).text = f"{Mno_warm_up}"

    doc.add_paragraph('''ИТОГО ВЫБРОСЫ''')

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    table.cell(0,0).text = "Код"
    table.cell(0,1).text = "Наименование ЗВ"
    table.cell(0,2).text = "Выброс г/с"
    table.cell(0,3).text = "Выброс т/год"


    table.cell(1,0).text = "0301"
    table.cell(1,1).text = "Азота диоксид"
    table.cell(1,2).text = f"{Gno2_warm_up}"
    table.cell(1,3).text = f"{Mno2_warm_up}"

    table.cell(2,0).text = "0304"
    table.cell(2,1).text = "Азот (II) оксид"
    table.cell(2,2).text = f"{Gno_warm_up}"
    table.cell(2,3).text = f"{Mno_warm_up}"

    table.cell(3,0).text = "0330"
    table.cell(3,1).text = "Сера диоксид"
    table.cell(3,2).text = f"{SOgik}"
    table.cell(3,3).text = f"{SO_mik}"

    table.cell(4,0).text = "0337"
    table.cell(4,1).text = "Углерода оксид"
    table.cell(4,2).text = f"{COgik_warmup}"
    table.cell(4,3).text = f"{CO_mik_warmup}"

    table.cell(5,0).text = "2704"
    table.cell(5,1).text = "Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(5,2).text = f"{petrol_gik_warmup}"
    table.cell(5,3).text = f"{petrol_mik_warmup}"

    doc.add_paragraph('''Максимально-разовые выбросы достигнуты в переходный период''')
    doc.save('prilozhenie/prilozhenie2/station.docx')

# generate_station_desription()
