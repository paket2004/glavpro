from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_page_break_doc():
    temp_doc = Document()
    paragraph = temp_doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return temp_doc

def combine_all_docx(filename_master, files_list):
    master = Document_compose(filename_master)
    composer = Composer(master)
    composer.append(create_page_break_doc())

    for i, file_path in enumerate(files_list):
        doc_temp = Document_compose(file_path)
        composer.append(doc_temp)
        if i < len(files_list) - 1:
            composer.append(create_page_break_doc())

    composer.save("prilozhenie\prilozhenie2\car_report.docx")



files_to_merge = [
    "prilozhenie/prilozhenie2/moving.docx",
]
# filename_master="report.docx"
filename_master="prilozhenie/prilozhenie2/station.docx"
combine_all_docx(filename_master,files_to_merge)