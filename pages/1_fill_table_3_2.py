import streamlit as st
import pandas as pd
from docx import Document
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.title("Редактируемая таблица источников выбросов")

columns = [
    "№ ИЗАВ", "Тип ИЗАВ", "Наименование ИЗАВ", "Число ИЗАВ, объединенных под одним номером",
    "Высота источника, м", "Диаметр устья (круглое), м", "Длина устья (прямоугольное), м", 
    "Ширина устья (прямоугольное), м", "Координаты X1", "Координаты Y1", 
    "Координаты X2", "Координаты Y2", "Ширина площадочного источника, м",
    "№ режима (стадии) выброса", "Скорость выхода ГВС, м/с (фактическая/осреднённая)",
    "Вертикальная составляющая осреднённой скорости выхода ГВС, м/с",
    "Объём (расход ГВС), м³/c (при фактических условиях) осреднённый",
    "Температура ГВС °C осреднённая", "Плотность ГВС, кг/м3",
    "КОД ЗВ", "Наименование ЗВ", "Концентрация мг/м3", "Мощность выброса, г/с",
    "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год",
    "Итого за год выброс веществ источником, т/год", "Примечание"
]



default_values = {
    "№ ИЗАВ": ["0001", "0002", "6003"],
    "Тип ИЗАВ": ["Организованный, точечный", "Организованный, точечный", "Неорганизованный, площадной"],
    "Наименование ИЗАВ": ["Труба котельной", "Труба продувочной свечи", "Открытая стоянка"],
    "Число ИЗАВ, объединенных под одним номером": ["1", "1", "1"],
    "Высота источника, м": [6.0, 6.0, 2.0],
    "Диаметр устья (круглое), м": [0.47, 0.032, ""],
    "Длина устья (прямоугольное), м": ["", "", ""],
    "Ширина устья (прямоугольное), м": ["", "", ""],
    "Координаты X1": ["174", "156", "168"],
    "Координаты Y1": ["174", "205", "109"],
    "Координаты X2": ["", "", "176"],
    "Координаты Y2": ["", "", "109"],
    "Ширина площадочного источника, м": ["", "", 3.0],
    "№ режима (стадии) выброса": ["1", "1", "1"],
    "Скорость выхода ГВС, м/с (фактическая/осреднённая)": [7, 5, ""],
    "Вертикальная составляющая осреднённой скорости выхода ГВС, м/с": ["", "", ""],
    "Объём (расход ГВС), м³/c (при фактических условиях) осреднённый": [1.2, 0.004, ""],
    "Температура ГВС °C осреднённая": [110.0, 20.0, ""],
    "Плотность ГВС, кг/м3": ["", "", ""],
    "КОД ЗВ": [["0301"], ["0304", "0337", "0703", "0410", "1728"], ["0301", "0304", "0330", "0337", "2704"]],
    "Наименование ЗВ": [
        ["Азота диоксид"], 
        ["Азот (II) оксид", "Углерода оксид", "Бенз/а/пирен", "Метан"], 
        ["Этантиол", "Азота диоксид", "Азот (II) оксид", "Сера диоксид", "Углерода оксид", "Бензин (нефтяной, малосернистый) в пересчете на углерод"]
    ],
    "Концентрация мг/м3": ["2.43", "0.35\n 7.51\n 0.000005","253.6\n 0.001"],
    "Мощность выброса, г/с": [
        "0.0021", 
        "0.0003\n 0.0065\n 0.000000004\n 0.00095\n 0.0000000042", 
        "0.000264\n 0.0000429\n 0.00006705\n 0.033398\n 0.0036225"
    ],
    "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год": [
        "0.0325", 
        "0.0053\n 0.0001\n 0.0000001\n 0.000014\n 0.00000000006", 
        "0.000065488\n 0.0000106418\n 0.000016922\n 0.008114\n 0.000865"
    ],
    "Итого за год выброс веществ источником, т/год": [
        "0.0325", 
        "0.0053\n 0.0001\n 0.0000001\n 0.000014\n 0.00000000006", 
        "0.000065488\n 0.0000106418\n 0.000016922\n 0.008114\n 0.000865"
    ],
    "Примечание": ["", "", ""]
}

if 'df_emissions' not in st.session_state:
    st.session_state.df_emissions = pd.DataFrame(default_values)

if st.button("Сбросить к значениям по умолчанию"):
    st.session_state.df_emissions = pd.DataFrame(default_values)
    st.rerun()

st.write("Введите данные об источниках выбросов:")
edited_df = st.data_editor(
    st.session_state.df_emissions, 
    num_rows="dynamic",
    column_config={
        "Высота источника, м": st.column_config.NumberColumn(format="%.2f"),
        "Диаметр устья (круглое), м": st.column_config.NumberColumn(format="%.2f"),
        "Длина устья (прямоугольное), м": st.column_config.NumberColumn(format="%.2f"),
        "Ширина устья (прямоугольное), м": st.column_config.NumberColumn(format="%.2f"),
        "Скорость выхода ГВС, м/с (фактическая/осреднённая)": st.column_config.NumberColumn(format="%.2f"),
        "Температура ГВС °C осреднённая": st.column_config.NumberColumn(format="%.1f"),
        "Плотность ГВС, кг/м3": st.column_config.NumberColumn(format="%.2f"),
        "Концентрация мг/м3": st.column_config.TextColumn(),
        "Мощность выброса, г/с": st.column_config.TextColumn(),
        "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год": st.column_config.TextColumn(),
        "Итого за год выброс веществ источником, т/год": st.column_config.TextColumn()
    },
    hide_index=True
)

# Сохраняем изменения в session_state
st.session_state.df_emissions1 = edited_df

# Кнопка для вывода данных
if st.button("Показать введённые данные"):
    st.write("Введённые данные об источниках выбросов:")
    st.dataframe(st.session_state.df_emissions1)

# Кнопка для сохранения таблицы в Word
if st.button("Сохранить таблицу в Word"):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(4)
    heading = doc.add_heading(level=1)
    run = heading.add_run("Стационарные источники выбросов загрязняющих веществ")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(13.5)
    section.page_height = Inches(10)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
     
    for index, row in st.session_state.df_emissions1.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            cell_value = row[col]
            
            if isinstance(cell_value, list):
                row_cells[i].text = "\n".join(map(str, cell_value))
            elif pd.isna(cell_value):
                row_cells[i].text = ""
            else:
                row_cells[i].text = str(cell_value)
    font_size = Pt(4)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = font_size
    for row in table.rows:
        row.cells[0].width = Inches(1.0)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(1.0)
    doc.save("calculations/tables/razdel3/3_2.docx")
    st.success("Таблица успешно сохранена в файл emission_sources_table.docx")

# Дополнительные элементы интерфейса
st.sidebar.header("Параметры таблицы")
st.sidebar.info("""
Используйте эту форму для ввода данных об источниках выбросов.
- Добавляйте новые строки с помощью кнопки '+'
- Редактируйте ячейки напрямую
- Числовые поля форматируются автоматически
- Для сброса к значениям по умолчанию нажмите соответствующую кнопку
""")