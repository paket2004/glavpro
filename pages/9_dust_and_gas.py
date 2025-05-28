import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document()
st.title("Характеристика пылегазоочистного оборудования и оценка его эффективности")

pgoy = st.text_input("Введите ПГОУ, установленное на вашем предприятии. Если такого нет, оставьте пустую строку")

if st.button("Отправить"):
    if pgoy:
        pass
    else:
        # 1.2.3 Характеристика пылегазоочистного оборудования и оценка его эффективности
        # Add heading (Times New Roman 8pt, bold, centered)
        heading = doc.add_paragraph()
        heading_run = heading.add_run('1.2.3 Характеристика пылегазоочистного оборудования и оценка его эффективности')
        heading_run.font.name = 'Times New Roman'
        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # For Cyrillic support
        heading_run.font.size = Pt(8)
        heading_run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add main text (Times New Roman 6pt)
        text_paragraph = doc.add_paragraph()
        text_run = text_paragraph.add_run("Пылегазоочистное оборудование (ПГОУ) на объекте негативного воздействия отсутствуют.")
        text_run.font.name = 'Times New Roman'
        text_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # For Cyrillic support
        text_run.font.size = Pt(6)
        
        doc.save("object_property/dust_and_gas/dust_and_gus_info.docx")
        st.success("Документ сохранён")