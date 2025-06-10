import os
import sys
sys.path.append(os.getcwd())
import streamlit as st
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import sys
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from gross_emissions.gross_emission_lists_for_different_source import car, boiler_pipe, blow_off_plug
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Настройка вкладок
tab1, tab2 = st.tabs(["Информация о компании", "Источники выбросов"])

with tab1:
    st.title("Информация о компании")

    columns_name = ["Наименование данных", "На момент разработки отчета инвентаризации"]

    if 'company_info_df' not in st.session_state:
        rows = [
            "Полное наименование предприятия",
            "Сокращенное наименование предприятия",
            "Адрес юридического лица",
            "Адреса осуществления деятельности", 
            "Фактический адрес площадки", 
            "ИНН", "ОГРН", "КПП", "ОКФС", "ОКОПФ", "ОКОГУ", "ОКТМО", "ОКПО","ОКВЭД",
            "Коды, присвоенные при постановке на государственный учет ОНВ", 
            "Директор",
            "Должностные лица, ответственные за проведение инвентаризации выбросов МБОУ «Мирошкинская СОШ»", 
            "Краткая характеристика местности, прилегающей к объекту ОНВ",
            "Размеры и границы санитарно-защитной зоны"
        ]
        
        data = {
            "Наименование данных": rows,
            "На момент разработки отчета инвентаризации": [""] * len(rows),
        }
        st.session_state.company_info_df = pd.DataFrame(data)

    st.write("Введите данные в таблицу:")
    edited_company_df = st.data_editor(st.session_state.company_info_df, num_rows="dynamic")
    st.session_state.company_info_df1 = edited_company_df

    if st.button("Показать введённые данные о компании"):
        st.write(st.session_state.company_info_df1)

with tab2:
    st.title("Информация об организации и источниках выбросов")

    st.subheader("Общая информация об организации")
    organization_info = {
        "Численность рабочих по штатному": st.text_input("Численность рабочих по штатному", "21 человек"),
        "Среднесписочная численность": st.text_input("Среднесписочная численность", "17 человек"),
        "Дата ввода объекта в эксплуатацию": st.text_input("Дата ввода объекта в эксплуатацию", "2001 г. (Котельная)"),
        "Адрес расположения площадки": st.text_input("Адрес расположения площадки", "461987, Оренбургская область, Первомайский район, с. Мирошкино, ул. Центральная, 27"),
        "Земельный участок": st.text_input("Земельный участок", "56:22:0501001:4 площадью 15 550 м²"),
        "Категория земель": st.text_input("Категория земель", "Земли населенных пунктов"),
        "Виды разрешённого использования": st.text_input("Виды разрешённого использования", "для размещения административного здания и производственных помещений"),
        "Арендаторы, осуществляющие негативное воздействие на окружающую среду": st.text_input("Арендаторы, осуществляющие негативное воздействие на окружающую среду", "отсутствуют"),
        "Водоснабжение": st.text_input("Водоснабжение", "централизованное"),
        "Водоотведение": st.text_input("Водоотведение", "выгребная яма"),
        "Теплоснабжение": st.text_input("Теплоснабжение", "от собственной котельной"),
        "Перспективы развития": st.text_input("Перспективы развития", "в ближайшие 7 лет не запланированы"),
        "На производственной площадке расположены следующие источники загрязнения атмосферы:": st.text_input("На производственной площадке расположены следующие источники загрязнения атмосферы:", "")
    }

    # Инициализация состояния для источников
    if "sources" not in st.session_state:
        st.session_state.sources = []
    if "next_source_id" not in st.session_state:
        st.session_state.next_source_id = 1

    st.subheader("Добавить источник")
    if st.button("Добавить источник"):
        st.session_state.sources.append({"id": st.session_state.next_source_id, "type": "Организованный", "data": {}})
        st.session_state.next_source_id += 1

    for source in st.session_state.sources:
        st.subheader(f"Источник №{source['id']}")
        source_type = st.selectbox(f"Тип источника для источника №{source['id']}", ["", "Труба котельной", "Продувочная свеча", "Открытая стоянка"], key=f"type_{source['id']}")

        if source_type == "Труба котельной":
            source["data"]["Номер источника"] = st.text_input(f"Номер источника для источника №{source['id']}", f"ИЗАВ №{source['id']:04d} – Труба котельной")
            source["data"]["ИВ (001)"] = st.text_input("Название источника для заполнения ИВ", f"Котлы водогрейные КЧМ -5", key=f"iv_{source['id']}")
            source["data"]["Тип котла"] = st.text_input(f"Тип котла для источника №{source['id']}", "водогрейный КЧМ -5")
            source["data"]["Мощность котла"] = st.text_input(f"Мощность котла для источника №{source['id']}", "21 кВт")
            source["data"]["Количество котлов в работе"] = st.text_input(f"Количество котлов в работе для источника №{source['id']}", "2 в работе, 0 в резерве")
            source["data"]["Количество одновременно работающих котлов"] = st.text_input(f"Количество одновременно работающих котлов для источника №{source['id']}", "2 шт.")
            source["data"]["Вид топлива"] = st.text_input(f"Вид топлива для источника №{source['id']}", "газ горючий природный")
            source["data"]["Фактический расход топлива на один котёл"] = st.text_input(f"Фактический расход топлива на один котёл для источника №{source['id']}", "15 024 м³")
            source["data"]["Плотность газа"] = st.text_input(f"Плотность газа для источника №{source['id']}", "0,733 кг/м³")
            source["data"]["Резервное топливо"] = st.text_input(f"Резервное топливо для источника №{source['id']}", "отсутствует")
            source["data"]["Количество дней работы в год"] = st.text_input(f"Количество дней работы в год для источника №{source['id']}", "181 день")
            source["data"]["Время работы в сутки"] = st.text_input(f"Время работы в сутки для источника №{source['id']}", "24 часа/день")
            source["data"]["Высота дымовой трубы"] = st.text_input(f"Высота дымовой трубы для источника №{source['id']}", "6 м")
            source["data"]["Диаметр дымовой трубы"] = st.text_input(f"Диаметр дымовой трубы для источника №{source['id']}", "0,470 м")
            source["data"]["Загрязняющие вещества"] = st.text_input(f"Загрязняющие вещества для источника №{source['id']}", f"{boiler_pipe}")

        elif source_type == "Продувочная свеча":
            source["data"]["Номер источника"] = st.text_input(f"Номер источника для источника №{source['id']}", f"ИЗАВ №{source['id']:04d} – Продувочная свеча")
            source["data"]["ИВ (001)"] = st.text_input("Название источника", f"Продувочная свеча", key=f"iv_{source['id']}")
            source["data"]["Время работы"] = st.text_input(f"Время работы для источника №{source['id']}", "1,5 час/день, 1 дней/год")
            source["data"]["Длина продуваемого трубопровода"] = st.text_input(f"Длина продуваемого трубопровода для источника №{source['id']}", "5 м")
            source["data"]["Внутренний диаметр продуваемого трубопровода"] = st.text_input(f"Внутренний диаметр продуваемого трубопровода для источника №{source['id']}", "57 мм")
            source["data"]["Диаметр трубы продувочной свечи"] = st.text_input(f"Диаметр трубы продувочной свечи для источника №{source['id']}", "32 мм")
            source["data"]["Высота трубы продувочной свечи"] = st.text_input(f"Высота трубы продувочной свечи для источника №{source['id']}", "6 м")
            source["data"]["Загрязняющие вещества"] = st.text_input(f"Загрязняющие вещества для источника №{source['id']}", f"{blow_off_plug}")

        elif source_type == "Открытая стоянка":
            source["data"]["Номер источника"] = st.text_input(f"Номер источника для источника №{source['id']}", f"ИЗАВ №{source['id']:04d} – Открытая стоянка")
            source["data"]["ИВ (001)"] = st.text_input("Название", f"Открытая стоянка", key=f"iv_{source['id']}")
            source["data"]["Размеры стоянки"] = st.text_input(f"Размеры стоянки для источника №{source['id']}", "3 х 8 м")
            source["data"]["Марка оборудования"] = st.text_input(f"Марка оборудования для источника №{source['id']}", "ГАЗ 322171")
            source["data"]["Количество ТС"] = st.text_input(f"Количество ТС для источника №{source['id']}", "1")
            source["data"]["Тип ТС"] = st.text_input(f"Тип ТС для источника №{source['id']}", "автобус")
            source["data"]["Время работы"] = st.text_input(f"Время работы для источника №{source['id']}", "1,5 часа/день, 181 дней/год")
            source["data"]["Вид топлива"] = st.text_input(f"Вид топлива для источника №{source['id']}", "Бензин")
            source["data"]["Загрязняющие вещества"] = st.text_input(f"Загрязняющие вещества для источника №{source['id']}", f"{car}")

            st.subheader("Таблица для Открытой стоянки")
            parking_data = {
                "Марка оборудования": [source["data"]["Марка оборудования"]],
                "Количество ТС": [source["data"]["Количество ТС"]],
                "Тип ТС": [source["data"]["Тип ТС"]],
                "Время работы": [source["data"]["Время работы"]],
                "Вид топлива": [source["data"]["Вид топлива"]]
            }
            st.table(parking_data)

def generate_full_report():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    
    # Установка минимальных отступов между параграфами
    style.paragraph_format.space_after = Pt(0)  # Убираем отступ после абзаца
    style.paragraph_format.line_spacing = 1.0   # Одинарный межстрочный интервал
    
    # Первый заголовок
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('''
    1. СВЕДЕНИЯ О ХОЗЯЙСТВУЮЩЕМ СУБЪЕКТЕ, ОБЪЕКТЕ ОНВ, ЕГО ОТДЕЛЬНЫХ ТЕРРИТОРИЯХ И ПРОИЗВОДСТВЕННОЙ ДЕЯТЕЛЬНОСТИ, ВКЛЮЧАЯ СВЕДЕНИЯ О КОЛИЧЕСТВЕ, ХАРАКТЕРИСТИКАХ И ЭФФЕКТИВНОСТИ ГОУ''')
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True
    
    # Второй заголовок
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('''
    1.1 Сведения о хозяйствующем субъекте, объекте ОНВ''')
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True
    
    # Таблица с данными компании
    company_table = doc.add_table(rows=st.session_state.company_info_df1.shape[0]+1, cols=2)
    company_table.style = 'Table Grid'
    company_table.cell(0, 0).text = "Категория"
    company_table.cell(0, 1).text = "Значение"
    
    for i, row in st.session_state.company_info_df.iterrows():
        company_table.cell(i+1, 0).text = str(row["Наименование данных"])
        company_table.cell(i+1, 1).text = str(row["На момент разработки отчета инвентаризации"])
    
    # Третий заголовок
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('''
    1.2 Сведения о производственной деятельности, количестве, характеристиках и эффективности ГОУ''')
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True
    
    # Общая информация об организации - с минимальными отступами
    for key, value in organization_info.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key}: {value}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)  # Минимальный отступ
    
    # Раздел с источниками выбросов - компактный
    for source in st.session_state.sources:    
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"{source['data'].get('Номер источника', '')}")
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        run.bold = True
        
        # Центрируем ИВ (001) для всех типов источников
        if 'ИВ (001)' in source['data']:
            p = doc.add_paragraph()
            run = p.add_run(f"ИВ (001): {source['data']['ИВ (001)']}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_after = Pt(0)
        
        # Обработка открытой стоянки
        if "Открытая стоянка" == source['data'].get('ИВ (001)', ''):
            p = doc.add_paragraph()
            run = p.add_run(f"Размеры стоянки: {source['data'].get('Размеры стоянки', '')}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            
            table = doc.add_table(rows=2, cols=5)
            table.style = 'Table Grid'
            headers = ["Марка оборудования", "Кол-во", "Тип ТС", "Время работы", "Вид топлива"]
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
                table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(8)
            table.cell(1, 0).text = source["data"].get("Марка оборудования", "")
            table.cell(1, 1).text = source["data"].get("Количество ТС", "")
            table.cell(1, 2).text = source["data"].get("Тип ТС", "")
            table.cell(1, 3).text = source["data"].get("Время работы", "")
            table.cell(1, 4).text = source["data"].get("Вид топлива", "")
            
            # Уменьшаем шрифт в таблице
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        
        # Добавляем остальные данные источника с минимальными отступами
        for key, value in source["data"].items():
            if key not in ["Марка оборудования", "Количество ТС", "Тип ТС", "Время работы", 
                          "Вид топлива", "Размеры стоянки", "Номер источника", "ИВ (001)"]:
                p = doc.add_paragraph()
                run = p.add_run(f"{key}: {value}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8)
                p.paragraph_format.space_after = Pt(0)
    
    # Сохранение документа
    if not os.path.exists("reports"):
        os.makedirs("reports")
    path_to_save = os.path.join(os.getcwd(), "inventarization_description", "punkt1", "organization_sources_info.docx")
    doc.save(path_to_save)
    return path_to_save

# Кнопка для генерации полного отчета
if st.button("Сформировать полный отчет в DOCX"):
    file_path = generate_full_report()
    with open(file_path, "rb") as f:
        st.download_button(
            "Скачать полный отчет",
            data=f,
            file_name="inventarization_description\izav_info\organization_sources_info.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )