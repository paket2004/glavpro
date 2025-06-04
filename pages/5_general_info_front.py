import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
st.title("Добавление источников выбросов и сохранение в DOCX")


work_area = st.text_input("Введите основной вид деятельности вашего предприятия")
organization_name = st.text_input("Введите название вашей организации")
# Инициализация состояния
if "sources_general_info" not in st.session_state:
    st.session_state.sources_general_info = []
if "next_org" not in st.session_state:
    st.session_state.next_org = 1  # 0001, 0002...
if "next_unorg" not in st.session_state:
    st.session_state.next_unorg = 6001  # 6001, 6002...

# Форма для добавления источников
st.subheader("Добавить новый источник выбросов")

category = st.radio("Тип источника", ["Организованный", "Неорганизованный"])

if category == "Организованный":
    number = f"{st.session_state.next_org:04d}"  # 0001, 0002...
else:
    number = str(st.session_state.next_unorg)  # 6001, 6002...

st.write(f"**Номер источника:** {number}")
name = st.text_input("Название источника")

if st.button("Добавить источник"):
    if name:
        st.session_state.sources_general_info.append({"number": number, "name": name, "category": category})
        if category == "Организованный":
            st.session_state.next_org += 1
        else:
            st.session_state.next_unorg += 1
    else:
        st.warning("Введите название источника!")

# Вывод списка источников
if st.session_state.sources_general_info:
    st.subheader("Список источников выбросов:")
    
    organized = [s for s in st.session_state.sources_general_info if s["category"] == "Организованный"]
    unorganized = [s for s in st.session_state.sources_general_info if s["category"] == "Неорганизованный"]

    st.markdown(f"**Всего на предприятии {len(st.session_state.sources_general_info)} источника выбросов:**")

    if organized:
        st.markdown(f"*Организованные источники ({len(organized)} шт.):*")
        for s in organized:
            st.markdown(f"Источник №{s['number']} – {s['name']}")

    if unorganized:
        st.markdown(f"*Неорганизованные источники ({len(unorganized)} шт.):*")
        for s in unorganized:
            st.markdown(f"Источник №{s['number']} – {s['name']}")

# Функция для генерации DOCX
def generate(sources):
    doc = Document()
    heading = doc.add_paragraph()
    heading_run = heading.add_run("2. ОПИСАНИЕ ПРОВЕДЕННЫХ РАБОТ ПО ИНВЕНТАРИЗАЦИИ С УКАЗАНИЕМ НОРМАТИВНО-МЕТОДИЧЕСКИХ ДОКУМЕНТОВ И ПЕРЕЧНЯ ИСПОЛЬЗОВАННЫХ МЕТОДИК ВЫПОЛНЕНИЯ ИЗМЕРЕНИЙ ЗАГРЯЗНЯЮЩИХ ВЕЩЕСТВ И РАСЧЁТНОГО ОПРЕДЕЛЕНИЯ ВЫБРОСОВ")
    heading_run.font.name = 'Times New Roman'
    heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # For Cyrillic support
    heading_run.font.size = Pt(8)
    heading_run.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        f"  Основной производственной деятельностью объекта негативного воздействия {organization_name} является {work_area}.\n"
        "   Всего на предприятии {} источника выбросов загрязняющих веществ в атмосферу из них:".format(len(sources))
    )

    organized = [s for s in sources if s["category"] == "Организованный"]
    unorganized = [s for s in sources if s["category"] == "Неорганизованный"]

    if organized:
        doc.add_paragraph(f"Организованные источники ({len(organized)} шт.):")
        for s in organized:
            doc.add_paragraph(f"Источник №{s['number']} – {s['name']}")

    if unorganized:
        doc.add_paragraph(f"Неорганизованные источники ({len(unorganized)} шт.):")
        for s in unorganized:
            doc.add_paragraph(f"Источник №{s['number']} – {s['name']}")
    doc.save("inventarization_description\izav_info\description.docx")
    return "inventarization_description\izav_info\description.docx"

# Кнопка для сохранения в DOCX
if st.button("Сохранить в DOCX"):
    if st.session_state.sources_general_info:
        file_path = generate(st.session_state.sources_general_info)
        with open(file_path, "rb") as f:
            st.download_button("Скачать DOCX", f, file_name="Источники_выбросов.docx")
    else:
        st.warning("Добавьте хотя бы один источник перед сохранением!")

# Кнопка для очистки списка
if st.button("Очистить список"):
    st.session_state.sources_general_info = []
    st.session_state.next_org = 1
    st.session_state.next_unorg = 6001
    st.experimental_rerun()
