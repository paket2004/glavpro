import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document()
st.title("Сведения о проведённых инвентаризациях")

history = st.text_input("Введите информацию о проведённых инвентаризациях (если их не было, оставьте поле пустым)")
org_name = st.text_input("Введите название вашей организации (сокращённое)", value="МБОУ «Мирошкинская СОШ»")

if st.button("Отправить"):
    if history:
        pass
    else:
        # Add heading (Times New Roman 8pt, bold, centered)
        heading = doc.add_paragraph()
        heading_run = heading.add_run("1.3 Сведения о результатах предыдущей инвентаризации")
        heading_run.font.name = 'Times New Roman'
        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # For Cyrillic support
        heading_run.font.size = Pt(8)
        heading_run.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add main text (Times New Roman 6pt)
        text_paragraph = doc.add_paragraph()
        text_run = text_paragraph.add_run(
            f"Ранее инвентаризации стационарных источников и выбросов вредных веществ в атмосферный воздух для {org_name} не проводилась."
        )
        text_run.font.name = 'Times New Roman'
        text_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # For Cyrillic support
        text_run.font.size = Pt(6)
        
        doc.save("object_property/previous_inventarization/prev_inv_info.docx")
        st.success("Документ сохранён")