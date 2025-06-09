import openai
from openai import OpenAI
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.shared import Pt, RGBColor
from docx import Document
import re
from termins_and_short.generate_short import shorcuts_list
import json
# from generate_short import shorcuts_list

openai.api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI()

def add_shortcuts_to_doc(doc, shortcuts_cur, shorcuts_list):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('''
    Принятые в отчете по инвентаризации стационарных источников и выбросов вредных (загрязняющих) веществ в атмосферный воздух сокращения:''')
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.bold = True
    for shorcut in shortcuts_cur:
        if shorcut in shorcuts_list.keys():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            run = p.add_run(f"{shorcut} – {shorcuts_list[shorcut]}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(6)
        else:
            print("Нет такого сокращения")
    


def generate_response(termins: str, shortcuts: dict):
    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant and your aim is to write an inventarization report for the company. Do not write abstract text, it should be very human written text."},
            {
                "role": "user",
                "content": f"""Дай определения на русском языке для следующих терминов: {termins}.
Ответ верни строго в формате JSON, где:
- ключ — это термин (включая всё, что в скобках, с запятыми и т.п.)
- значение — определение этого термина.

Пример:
{{
  "история": "наука, изучающая человеческое прошлое...",
  "Очистка газа (воздуха), пылегазоочистка": "Процессы удаления загрязнений..."
}}

Не добавляй никакого текста до или после JSON. Только JSON.
"""
            }
        ]
    )

    response = completion.choices[0].message.content.strip()

    try:
        results = json.loads(response)
    except json.JSONDecodeError as e:
        print("❌ Ошибка JSON:", e)
        print("Ответ от модели:\n", response)
        results = {}

    if not results:
        raise ValueError("Ответ от OpenAI пустой или нераспарсенный — проверь формат или термины.")

    # Генерация DOCX
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'

    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('Основные термины, используемые в проекте')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)

    table = doc.add_table(rows=len(results), cols=2)
    table.style = 'Table Grid'

    for i, (key, value) in enumerate(results.items()):
        table.cell(i, 0).text = str(key)
        table.cell(i, 1).text = str(value)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)

    if shortcuts:
        add_shortcuts_to_doc(doc, shortcuts, shorcuts_list)

    doc.save('termins_and_short/dictionary_table.docx')
    

