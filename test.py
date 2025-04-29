from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
def add_page_break(doc):
    """Inserts a page break into the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
def merge_documents_with_styles(master_path, files_to_merge, output_path):
    # Open master document
    master = Document(master_path)
    
    for file_path in files_to_merge:
        master = Document(master_path)
    
    for i, file_path in enumerate(files_to_merge):
        if i > 0:
            add_page_break(master)
        
        src_doc = Document(file_path)
        
        for element in src_doc.element.body:
            if element.tag.endswith('sectPr'):
                continue
            master.element.body.append(element)
    
    master.save(output_path)

files_to_merge = [
    "termins_and_short/dictionary_table.docx",
    "introduction/introduction.docx",
    "inventarization_description\punkt1\organization_sources_info.docx",
    "object_property\sanitary_zone\sanitary_zone.docx",
    r"inventarization_description\izav_info\razdel2.docx",
    r"calculations\tables\razdel3\razdel3.docx"
]
# Usage
merge_documents_with_styles(
    "content/Содержание_отчета.docx",
    files_to_merge,
    "combined_with_styles.docx"
)