import streamlit as st
import pandas as pd
from docx import Document
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
# Заголовок приложения
st.title("Редактируемая таблица источников выбросов")

# Задаём колонки таблицы на основе шаблона
columns = [
    "№ ИЗАВ", "Тип ИЗАВ", "Наименование ИЗАВ", "Число ИЗАВ, объединенных под одним номером",
    "Высота источника, м", "Диаметр устья (круглое), м", "Длина устья (прямоугольное), м", 
    "Ширина устья (прямоугольное), м", "Координаты X1", "Координаты Y1", 
    "Координаты X2", "Координаты Y2", "Ширина площадочного источника, м",
    "№ режима (стадии) выброса", "Скорость выхода ГВС, м/с (фактическая/осреднённая)",
    "Вертикальная составляющая осреднённой скорости выхода ГВС, м/с",
    "Объём (расход ГВС), м³/c (при фактических условиях) осреднённый",
    "Температура ГВС °C осреднённая", "Плотность ГВС, кг/м3",
    "КОД ЗВ", "Наименование ЗВ", "Концентрация мг/м3", "Мощность выброса, г/с",
    "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год",
    "Итого за год выброс веществ источником, т/год", "Примечание"
]

# Default значения для таблицы
default_values = {
    "№ ИЗАВ": ["001", "002"],
    "Тип ИЗАВ": ["Труба", "Площадка"],
    "Наименование ИЗАВ": ["Котельная №1", "Автостоянка"],
    "Число ИЗАВ, объединенных под одним номером": ["1", "1"],
    "Высота источника, м": [25.0, 5.0],
    "Диаметр устья (круглое), м": [1.5, ""],
    "Длина устья (прямоугольное), м": ["", 10.0],
    "Ширина устья (прямоугольное), м": ["", 8.0],
    "Координаты X1": ["123.456", "123.458"],
    "Координаты Y1": ["456.789", "456.791"],
    "Координаты X2": ["", ""],
    "Координаты Y2": ["", ""],
    "Ширина площадочного источника, м": ["", 15.0],
    "№ режима (стадии) выброса": ["1", "1"],
    "Скорость выхода ГВС, м/с (фактическая/осреднённая)": [12.5, 0.5],
    "Вертикальная составляющая осреднённой скорости выхода ГВС, м/с": [10.0, ""],
    "Объём (расход ГВС), м³/c (при фактических условиях) осреднённый": [5.6, 2.3],
    "Температура ГВС °C осреднённая": [120.0, 25.0],
    "Плотность ГВС, кг/м3": [1.2, 1.2],
    "КОД ЗВ": ["0337", "2704"],
    "Наименование ЗВ": ["Углерод оксид", "Пыль неорганическая"],
    "Концентрация мг/м3": [125.4567, 50.1234],
    "Мощность выброса, г/с": [0.1234, 0.0567],
    "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год": [0.001234, 0.000567],
    "Итого за год выброс веществ источником, т/год": [0.001234, 0.000567],
    "Примечание": ["Основной источник", "Сезонная работа"]
}

# Инициализация DataFrame с default значениями
if 'df_emissions' not in st.session_state:
    st.session_state.df_emissions = pd.DataFrame(default_values)

# Кнопка для сброса к default значениям
if st.button("Сбросить к значениям по умолчанию"):
    st.session_state.df_emissions = pd.DataFrame(default_values)
    st.rerun()

# Отображаем редактируемую таблицу с возможностью добавления строк
st.write("Введите данные об источниках выбросов:")
edited_df = st.data_editor(
    st.session_state.df_emissions, 
    num_rows="dynamic",
    column_config={
        "Высота источника, м": st.column_config.NumberColumn(format="%.2f"),
        "Диаметр устья (круглое), м": st.column_config.NumberColumn(format="%.2f"),
        "Длина устья (прямоугольное), м": st.column_config.NumberColumn(format="%.2f"),
        "Ширина устья (прямоугольное), м": st.column_config.NumberColumn(format="%.2f"),
        "Скорость выхода ГВС, м/с (фактическая/осреднённая)": st.column_config.NumberColumn(format="%.2f"),
        "Температура ГВС °C осреднённая": st.column_config.NumberColumn(format="%.1f"),
        "Плотность ГВС, кг/м3": st.column_config.NumberColumn(format="%.2f"),
        "Концентрация мг/м3": st.column_config.NumberColumn(format="%.4f"),
        "Мощность выброса, г/с": st.column_config.NumberColumn(format="%.4f"),
        "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год": st.column_config.NumberColumn(format="%.6f"),
        "Итого за год выброс веществ источником, т/год": st.column_config.NumberColumn(format="%.6f")
    },
    hide_index=True
)

# Сохраняем изменения в session_state
st.session_state.df_emissions = edited_df

# Кнопка для вывода данных
if st.button("Показать введённые данные"):
    st.write("Введённые данные об источниках выбросов:")
    st.dataframe(st.session_state.df_emissions)

# Кнопка для сохранения таблицы в Word
if st.button("Сохранить таблицу в Word"):
    # Создаём новый документ Word
    doc = Document()
    
    # Добавляем заголовок
    doc.add_heading("Таблица источников выбросов", level=1)
    section = doc.sections[0]
    section._sectPr.xpath('./w:pgSz')[0].set(qn('w:orient'), 'landscape')

    # Adjust page width and height for landscape orientation
    section.page_width = Inches(11)  # Landscape width (11 inches)
    section.page_height = Inches(8.5)  # Landscape height (8.5 inches)

    # Создаём таблицу в Word
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    # Добавляем заголовки столбцов
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    
    # Добавляем данные из DataFrame в таблицу Word
    for index, row in st.session_state.df_emissions.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            row_cells[i].text = str(row[col] if pd.notna(row[col]) else "")
    
    # Сохраняем документ
    doc.save("calculations/tables/razdel3/3_2.docx")
    st.success("Таблица успешно сохранена в файл emission_sources_table.docx")

# Дополнительные элементы интерфейса
st.sidebar.header("Параметры таблицы")
st.sidebar.info("""
Используйте эту форму для ввода данных об источниках выбросов.
- Добавляйте новые строки с помощью кнопки '+'
- Редактируйте ячейки напрямую
- Числовые поля форматируются автоматически
- Для сброса к значениям по умолчанию нажмите соответствующую кнопку
""")