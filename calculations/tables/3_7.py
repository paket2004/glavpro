from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
# Создаем новый документ
doc = Document()
style = doc.styles['Normal']

style.font.name = 'Times New Roman'  # Название шрифта
style.font.size = Pt(6)  # Размер шрифта (12 пунктов)
# doc.add_heading('Результаты обследования установок очистки газа и условий их эксплуатации', 0)
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('Таблица № 3.7. Суммарные выбросы загрязняющих веществ в атмосферу, их очистка и утилизация (в целом по предприятию), т/год.')
run.font.color.rgb = RGBColor(0,0,0)
run.font.name = 'Times New Roman'  # This is the key line you're missing
run.font.size = Pt(8)

# Change the page orientation to landscape
section = doc.sections[0]
section._sectPr.xpath('./w:pgSz')[0].set(qn('w:orient'), 'landscape')

# Adjust page width and height for landscape orientation
section.page_width = Inches(11)  # Landscape width (11 inches)
section.page_height = Inches(8.5)  # Landscape height (8.5 inches)
table = doc.add_table(rows=4, cols=10)
table.style = 'Table Grid'

# Заголовки столбцов
headers = [
    "", "Загрязняющее вещество", "Количество загрязняющих веществ, отходящих от источников выделения",
    "", "Выбрасывается без очистки", "Поступает на очистку",  "", "", "Из поступивших на очистку",
    "Всего выброшено в атмосферный воздух"
]

# Объединяем ячейки для заголовков
table.cell(0, 0).merge(table.cell(0, 1))  # N ИЗАВ
table.cell(1,0).text = "Код"
table.cell(1,1).text = "Наименование"
table.cell(1, 0).merge(table.cell(2, 0))
table.cell(1, 1).merge(table.cell(2, 1))
table.cell(0, 2).merge(table.cell(2, 2))  # ТИП ИЗАВ
table.cell(0, 3).merge(table.cell(0, 4))  # НАИМАНОВАНИЕ ИЗАВ
table.cell(1,3).text = "Всего"
table.cell(1,4).text = "В том числе от организованных ИЗАВ"
table.cell(1, 3).merge(table.cell(2, 3))
table.cell(1, 4).merge(table.cell(2, 4))
table.cell(0, 5).merge(table.cell(2, 5))  # Число ИЗАВ, объединенных под одним номером
table.cell(0, 6).merge(table.cell(0, 8))
table.cell(1, 6).merge(table.cell(1, 7))
table.cell(1, 8).merge(table.cell(2, 8))
table.cell(1,6).text = "Уловлено и обезврежено"
table.cell(2,6).text = "Фактически"
table.cell(2,7).text = "Из них утилизировано"
table.cell(1,8).text = "Выброшено в атмосферный воздух"
table.cell(0, 9).merge(table.cell(2, 9)) 

# Заполняем заголовки
for i, header in enumerate(headers):
    table.cell(0, i).text = header

    # # Добавляем пустые строки для данных
    # for row in range(2, 3):  # Добавляем одну пустую строку для данных
    #     for col in range(18):
    #         table.cell(row, col).text = ""  # Оставляем ячейки пустыми
    # Сохраняем документ  
doc.save('calculations/tables/3_7.docx')