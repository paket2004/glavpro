from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Создаем новый документ
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(8)
doc.add_heading('Отчёт об инвентаризации выбросов', 0)

# Change the page orientation to landscape
section = doc.sections[0]
section._sectPr.xpath('./w:pgSz')[0].set(qn('w:orient'), 'landscape')

# Adjust page width and height for landscape orientation
section.page_width = Inches(11)  # Landscape width (11 inches)
section.page_height = Inches(8.5)  # Landscape height (8.5 inches)
table = doc.add_table(rows=4, cols=18)
table.style = 'Table Grid'

# Заголовки столбцов
headers = [
    "№ цеха", "Наименование цеха", "№ участка", "Наименование участка",
    "Номер источника выделения (ИВ)", "Наименование источника выделения (ИВ)",
    "Характеристика нестационарности разработки ИВ (№ режима нестационарности)",
    "", "Время работы ИВ с учетом нестационарности",  "Количество ИВ под одним номером",
    "", "Загрязняющее вещество",  "", "", "Количество ЗВ, отходящих от ИВ",
    "Инвентарный № газоочистного оборудования - установок очистки газа (если проводится очистка)",
    "Номер ИЗАВ, в который поступают загрязняющие вещества от источника выделения",
    "Примечание"
]

table.cell(0, 0).merge(table.cell(2, 0))
table.cell(0, 1).merge(table.cell(2, 1))
table.cell(0, 2).merge(table.cell(2, 2))
table.cell(0, 3).merge(table.cell(2, 3))
table.cell(0, 4).merge(table.cell(2, 4))
table.cell(0, 5).merge(table.cell(2, 5))
table.cell(0, 6).merge(table.cell(2, 6))

merged_cell = table.cell(0, 7).merge(table.cell(0, 8))
table.cell(1, 7).text = "В сутки, час/сутки"
table.cell(1, 8).text = "Всего за год, часов"
table.cell(1, 7).merge(table.cell(2, 7)) 
table.cell(1, 8).merge(table.cell(2, 8)) 

table.cell(0, 9).merge(table.cell(2, 9))

table.cell(0, 10).merge(table.cell(0, 11))
table.cell(1,10).text = "Код"
table.cell(1,11).text = "Наименование"
table.cell(1, 10).merge(table.cell(2, 10)) 
table.cell(1, 11).merge(table.cell(2, 11)) 

table.cell(0, 12).merge(table.cell(0, 14))
table.cell(1, 12).merge(table.cell(1, 13))
table.cell(1,13).text = "При учете нестационарности"
table.cell(2,12).text = "г/с"
table.cell(2,13).text = "т/год"
table.cell(1,14).text = "Всего (тонн в год)"

table.cell(0, 15).merge(table.cell(2, 15))
table.cell(0, 16).merge(table.cell(2, 16))
table.cell(0, 17).merge(table.cell(2, 17))


table.columns[0].width = Inches(1.5)
table.columns[1].width = Inches(6)
table.columns[6].width = Inches(1.5)
# table.columns[3].width = Inches(1.5)
# table.columns[4].width = Inches(6)

# table.columns[5].width = Inches(1.5)
# table.columns[6].width = Inches(6)
# table.columns[7].width = Inches(1.5)
# table.columns[8].width = Inches(6)


for i, header in enumerate(headers):
    table.cell(0, i).text = header
doc.save('calculations/tables/3_1.docx')