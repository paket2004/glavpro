from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
# Создаем новый документ
doc = Document()
style = doc.styles['Normal']

style.font.name = 'Times New Roman'  # Название шрифта
style.font.size = Pt(6)  # Размер шрифта (12 пунктов)
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('Таблица № 3.8 Выбросы от передвижных ИЗАВ на 2024 год')
run.font.color.rgb = RGBColor(0,0,0)
run.font.name = 'Times New Roman'  # This is the key line you're missing
run.font.size = Pt(8)


# Change the page orientation to landscape
section = doc.sections[0]
section._sectPr.xpath('./w:pgSz')[0].set(qn('w:orient'), 'landscape')

# Adjust page width and height for landscape orientation
section.page_width = Inches(11)  # Landscape width (11 inches)
section.page_height = Inches(8.5)  # Landscape height (8.5 inches)
table = doc.add_table(rows=4, cols=11)
table.style = 'Table Grid'

# Заголовки столбцов
headers = [
    "№", "ИЗАВ, его вид (согласно п.5 настоящего порядка)", "Количество ИЗАВ каждого вида", "Скорость движения ИЗАВ по объекту ОНВ, (км/ч)",
    "Вид топлива", "Время работы за сезон, (ч)", "Время работы за год, (ч) ",  "", "", "Выброс загрязняющих веществ", "Ссылка на расчетную методику",
]

# Объединяем ячейки для заголовков
table.cell(0, 0).merge(table.cell(1, 0))  # N ИЗАВ
table.cell(0, 1).merge(table.cell(1, 1))  # ТИП ИЗАВ
table.cell(0, 2).merge(table.cell(1, 2))  # НАИМАНОВАНИЕ ИЗАВ
table.cell(0, 3).merge(table.cell(1, 3))  # Число ИЗАВ, объединенных под одним номером
table.cell(0, 4).merge(table.cell(1, 4)) 
table.cell(0, 5).merge(table.cell(1, 5))  # НАИМАНОВАНИЕ ИЗАВ
table.cell(0, 6).merge(table.cell(1, 6))  
table.cell(0, 7).merge(table.cell(0, 9)) 
table.cell(1,7).text = "Наименование загрязняющего вещества"
table.cell(1,8).text = "Выбросы ЗВ, макс,(г/с)"
table.cell(1,9).text = "Выбросы ЗВ, за год (т/год)"

table.cell(0, 10).merge(table.cell(1, 10))
# Заполняем заголовки
for i, header in enumerate(headers):
    table.cell(0, i).text = header

    # # Добавляем пустые строки для данных
    # for row in range(2, 3):  # Добавляем одну пустую строку для данных
    #     for col in range(18):
    #         table.cell(row, col).text = ""  # Оставляем ячейки пустыми
    # Сохраняем документ  
# calculations\tables\3_8.py
doc.save('calculations/tables/3_8.docx')