# ВЫБРОСЫ ОТ АВТОМОБИЛЯ ПОСЛЕ ЕГО ПРОГРЕВА!
from docx.shared import Pt 

from docx import Document
doc = Document("calculations/tables/3_1.docx")
def fill_table_3_1(sera_dioksid_sec, sera_dioksid_year, Mno_warm, SO_mik_warm, CO_mik_warm, petrol_mik_warm, Gno_warm, SO_gik_warm, COgik_warm, petrol_gik_warm, mco_sec, mco_year, MNO2_g, azot_2_oksid_year, mno2_t, mno_g, benzapiren_sec,benzapiren_year, methane_sec, methane_year,etantiol_sec, etantiol_year, Gno2_warm, Gno2_warm_up, Mno2_warm, Mno2_warm_up, SOgik, SO_mik, COgik_warmup, CO_mik_warmup, petrol_gik_warmup, petrol_mik_warmup, num_of_boilers, num_of_days_in_work,num_of_hours_in_day_in_work, num_of_days_in_work_parking, num_of_hours_in_day_in_work_parking, num_of_parkings, num_of_days_in_work_candle, num_of_hours_in_day_in_work_candle, num_of_candles, num_of_days_in_work_car, num_of_hours_in_day_in_work_car, car_amount):
    table = doc.tables[0]
    # Set font size for all cells (example: 10pt)
    font_size = Pt(4)  # Adjust size as needed
    
    table.cell(3,0).text = "01\n\n\n\n02"
    # Наименование цеха
    table.cell(3,1).text = "Котельная\n\n\n\nСтоянка транспорта"
    # Номер источника выделения (ИВ)
    table.cell(3,4).text =  "001\n\n001\n\n001\n\n\n\n002"
    # Наименование источника выделения (ИВ)
    table.cell(3,5).text = "Котлы водогрейные КЧМ -5 \n\n Продувочная свеча \n Открытая стоянка \n Движение и работа транспорта по территории (автобус)"
    # Характеристика нестационарнсти работы ИВ (№ режима нестационарности)
    table.cell(3,6).text =  "1\n\n1\n\n1\n\n\n\n1"

    table.cell(3,7).text = f"{num_of_boilers * num_of_hours_in_day_in_work}\n\n{num_of_candles * num_of_hours_in_day_in_work_candle}\n\n{num_of_parkings * num_of_hours_in_day_in_work_parking}\n\n{car_amount * num_of_hours_in_day_in_work_car}"
    
    table.cell(3,8).text = f"{num_of_boilers * num_of_hours_in_day_in_work * num_of_days_in_work}\n\n{num_of_candles * num_of_hours_in_day_in_work_candle * num_of_days_in_work_candle}\n\n{num_of_parkings * num_of_days_in_work_parking * num_of_hours_in_day_in_work_parking}\n\n{car_amount * num_of_hours_in_day_in_work_car * num_of_days_in_work_car}"

    table.cell(3,9).text = f"{num_of_boilers}\n\n{num_of_candles}\n\n{num_of_parkings}\n\n{car_amount}"
    # Код ЗВ
    table.cell(3,10).text = "0301\n\n0304\n\n0337\n\n0703\n\n0410\n\n1728\n\n0301\n\n0304\n\n0330\n\n0337\n\n2704\n\n0301\n\n0304\n\n0330\n\n0337\n\n2704\n\n0301\n\n0304\n\n0330\n\n0337\n\n2704"
    # Наименование ЗВ
    table.cell(3,11).text = "Азота диоксид\n\nАзот (II) оксид\n\nУглерода оксид\n\nБенз/а/пирен\n\nМетан\n\nЭтантиол\n\nАзота диоксид\n\nАзот (II) оксид\n\nСера диоксид\n\nУглерода оксид\n\nБензин (нефтяной, малосернистый) /в пересчете на углерод/\n\n Азота диоксид \n\n Азот (II) оксид \n\n Сера диоксид \n\n Углерода оксид \n\n Бензин (нефтяной,малосернистый) /в пересчете на углерод/"
    # Кол-во выделяемого ЗВ (в граммах/сек и т/год) 
    table.cell(3,12).text = f"{MNO2_g}\n\n{mno_g}\n\n{mco_sec}\n\n{benzapiren_sec}\n\n{methane_sec}\n\n{etantiol_sec}\n\n{Gno2_warm_up}\n\n{0.0000364}\n\n{SOgik}\n\n{COgik_warmup}\n\n{petrol_gik_warmup}\n\n{Gno2_warm}\n\n{Gno_warm}\n\n{SO_gik_warm}\n\n{COgik_warm}\n\n{petrol_gik_warm}"
    #                             +              +                    +             +                    +                  +               +                   -                +           -              +                    -           -                          +            +                +                  +                  +                 +                   +
    table.cell(3,13).text = f"{mno2_t}\n\n{azot_2_oksid_year}\n\n{mco_year}\n\n{benzapiren_year}\n\n{methane_year}\n\n{etantiol_year}\n\n{Mno2_warm_up}\n\n{"0.00000923"}\n\n{SO_mik}\n\n{CO_mik_warmup}\n\n {petrol_mik_warmup}\n\n {Mno2_warm} \n\n {Mno_warm} \n\n {SO_mik_warm} \n\n {CO_mik_warm}\n\n {petrol_mik_warm}"
    table.cell(3,14).text = f"{mno2_t}\n\n{azot_2_oksid_year}\n\n{mco_year}\n\n{benzapiren_year}\n\n{methane_year}\n\n{etantiol_year}\n\n{Mno2_warm_up}\n\n{"0.00000923"}\n\n{SO_mik}\n\n{CO_mik_warmup}\n\n {petrol_mik_warmup}\n\n {Mno2_warm} \n\n {Mno_warm} \n\n {SO_mik_warm} \n\n {CO_mik_warm}\n\n {petrol_mik_warm}"
    

    # Iterate through all cells in the table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = font_size
    
    doc.save("calculations/tables/razdel3/3_1.docx")