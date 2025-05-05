from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
# Создаем новый документ
doc = Document()
style = doc.styles['Normal']

style.font.name = 'Times New Roman'  # Название шрифта
style.font.size = Pt(6)  # Размер шрифта (12 пунктов)

heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('Стационарные источники выбросов загрязняющих веществ')
run.font.color.rgb = RGBColor(0,0,0)
run.font.name = 'Times New Roman'  # This is the key line you're missing
run.font.size = Pt(8)
# Change the page orientation to landscape
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
# Adjust page width and height for landscape orientation
section.page_width = Inches(11)  # Landscape width (11 inches)
section.page_height = Inches(8.5)  # Landscape height (8.5 inches)
table = doc.add_table(rows=4, cols=26)
table.style = 'Table Grid'

# Заголовки столбцов
headers = [
    "№ ИЗАВ", "Тип ИЗАВ", "Наименование ИЗАВ", "Число ИЗАВ, объединенных под одним номером",
    "Высота источника, м", "", "", "Размеры устья источника", "", "", "", "Координаты источника на карте-схеме",  "Ширина площадочного источника, м",
    "№ режима (стадии) выброса", "Скорость выхода ГВС, м/с, фактическая/осреднённая",  "Вертикальная составляющая осреднённой скорости выхода ГВС, м/с", 
    "Объём (расход ГВС), м^3/c (при фактических условиях) осреднённый", 
    "Температура ГВС С осреднённая", "Плотность ГВС, кг/м3", "", "", "", "",
    "ЗВ, выбрасываемые в атмосферный воздух (для каждого режима (стадии) выброса ЗВ)",
    "Итого за год выброс вещества источником, т/год",
    "Примечание"
]

# Объединяем ячейки для заголовков
table.cell(0, 0).merge(table.cell(2, 0))  # N ИЗАВ
table.cell(0, 1).merge(table.cell(2, 1))  # ТИП ИЗАВ
table.cell(0, 2).merge(table.cell(2, 2))  # НАИМАНОВАНИЕ ИЗАВ
table.cell(0, 3).merge(table.cell(2, 3))  # Число ИЗАВ, объединенных под одним номером
table.cell(0, 4).merge(table.cell(2, 4))  # Высота источника, м

table.cell(0, 5).merge(table.cell(0, 7))  # Количество ЗВ, отходящих от ИВ
table.cell(1, 5).merge(table.cell(1, 5))  # Круглое устье
table.cell(1, 5).text = "Круглое устье"
table.cell(2, 5).text = "диаметр, м"

table.cell(1, 6).merge(table.cell(1, 7))  # Прямоугольное устье
table.cell(1, 6).text = "Прямоугольное устье"
table.cell(2, 6).text = "Длина, м"  # Длина для прямоугольного устья
table.cell(2, 7).text = "Ширина, м"  # Ширина для прямоугольного устья

table.cell(0, 8).merge(table.cell(0, 11))  # Наименование источника выделения (ИВ)
table.cell(1, 8).merge(table.cell(1, 11))
table.cell(0, 8).merge(table.cell(1, 8))
table.cell(2, 8).text = "X1"  # Подзаголовок для первой ячейки
table.cell(2, 9).text = "Y1"  # Подзаголовок для второй ячейки
table.cell(2, 10).text = "X2"  # Подзаголовок для первой ячейки
table.cell(2, 11).text = "Y2"  # Подзаголовок для второй ячейки

table.cell(0, 12).merge(table.cell(2, 12)) 
table.cell(0, 13).merge(table.cell(2, 13)) 
table.cell(0, 14).merge(table.cell(2, 14)) 
table.cell(0, 15).merge(table.cell(2, 15)) 
table.cell(0, 16).merge(table.cell(2, 16))
table.cell(0, 17).merge(table.cell(2, 17)) 
table.cell(0, 18).merge(table.cell(2, 18))  

table.cell(0, 19).merge(table.cell(0, 23))
table.cell(1,19).text = "КОД"
table.cell(1,20).text = "Наименование"
table.cell(1,21).text = "Концентрация мг/м3"
table.cell(1,22).text = "Мощность выброса, г/с"
table.cell(1,23).text = "Суммарные годовые (валовые) выбросы режима (стадии) ИЗАВ, т/год"
table.cell(1, 19).merge(table.cell(2, 19))
table.cell(1, 20).merge(table.cell(2, 20))
table.cell(1, 21).merge(table.cell(2, 21))
table.cell(1, 22).merge(table.cell(2, 22))
table.cell(1, 23).merge(table.cell(2, 23))
table.cell(0, 24).merge(table.cell(2, 24))
table.cell(0, 25).merge(table.cell(2, 25))

# Заполняем заголовки
for i, header in enumerate(headers):
    table.cell(0, i).text = header

doc.save('calculations/tables/3_2.docx')