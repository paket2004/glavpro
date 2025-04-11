from docx import Document

doc = Document()
doc.add_table(rows=1, cols=5)
table = doc.tables[0]
table.style = 'Table Grid'
table.cell(0,0).text = "Марка оборудования"
table.cell(0,1).text = "Кол-во"
table.cell(0,2).text = "Тип ТС"
table.cell(0,3).text = "Время работы час/день, дней/год на территории предприятия"
table.cell(0,4).text = "Вид топлива"
table.add_row()
doc.save("calculations/tables/parking.docx")