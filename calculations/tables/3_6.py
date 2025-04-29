from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
# Создаем новый документ
doc = Document()
# style = doc.styles['Normal']
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(5)
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('''Результаты обследования установок очистки газа и условий их эксплуатации''')
run.font.name = 'Times New Roman'  # This is the key line you're missing
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0,0,0)
# style.font.name = 'Times New Roman'  # Название шрифта
# style.font.size = Pt(6)  # Размер шрифта (12 пунктов)
# # doc.add_heading('Результаты обследования установок очистки газа и условий их эксплуатации', 0)
# heading = doc.add_heading(level=1)
# heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
# run = heading.add_run('Результаты обследования установок очистки газа и условий их эксплуатации')
# run.font.color.rgb = RGBColor(0,0,0)
# run.font.name = 'Times New Roman'  # This is the key line you're missing
# run.font.size = Pt(8)


# Change the page orientation to landscape
section = doc.sections[0]
section._sectPr.xpath('./w:pgSz')[0].set(qn('w:orient'), 'landscape')

# Adjust page width and height for landscape orientation
section.page_width = Inches(11)  # Landscape width (11 inches)
section.page_height = Inches(8.5)  # Landscape height (8.5 inches)
table = doc.add_table(rows=3, cols=11)
table.style = 'Table Grid'

# Заголовки столбцов
headers = [
    "Результаты обследования установок очистки газа и условий их эксплуатации, № цеха ", "Наименование цеха", "№ участка", "Наименование источника выделения (выброса), его номер",
    "Наименование установок отчистки газа, его тип и марка (№ в реестре установок очистки газа на объекте ОНВ)", "Номер ИЗАВ, через который осуществляются выбросы после очистки", "", 
    "Эффективность (степень очистки) установок очистки газа, %", "Наименование и код ЗВ", "", "Коэффициент обеспеченности, %"
]

# Объединяем ячейки для заголовков
table.cell(0, 0).merge(table.cell(1, 0))  # N ИЗАВ
table.cell(0, 1).merge(table.cell(1, 1))  # ТИП ИЗАВ
table.cell(0, 2).merge(table.cell(1, 2))  # НАИМАНОВАНИЕ ИЗАВ
table.cell(0, 3).merge(table.cell(1, 3))  # Число ИЗАВ, объединенных под одним номером
table.cell(0, 4).merge(table.cell(1, 4))  # Высота источника, м
table.cell(0, 5).merge(table.cell(1, 5)) 
table.cell(0, 6).merge(table.cell(0, 7))  # Количество ЗВ, отходящих от ИВ
table.cell(1,6).text = "Проектный"
table.cell(1,7).text = "Фактический"

table.cell(0, 8).merge(table.cell(1, 8)) 

table.cell(0, 9).merge(table.cell(0, 10))  # Количество ЗВ, отходящих от ИВ
table.cell(1,9).text = "Нормативный"
table.cell(1,10).text = "Фактический"


# Заполняем заголовки
for i, header in enumerate(headers):
    table.cell(0, i).text = header
table.cell(2,0).merge(table.cell(2, 10))
# Выбираем ячейку (2, 0)
cell = table.cell(2, 0)

# Устанавливаем текст
cell.text = "Пылегазоочистное оборудование отсутствует!"

# Выравниваем текст по центру
for paragraph in cell.paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'calculations\tables\razdel3\3_6.docx')