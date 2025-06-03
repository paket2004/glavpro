import os
import sys
sys.path.append(os.getcwd())

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from calculations.tables.fill_table_3_1 import fill_table_3_1
from calculations.tables.fill_table_3_8 import fill_3_8_table

from calculations.gas.benzapiren import calculate_amount_of_benzapiren_sec, calculate_amount_of_benzapiren_year
from calculations.gas.nitric_ocide import calculate_MNO2_g, calculate_Mno2_t, calculate_MNO_g, calculate_MNO_t
from calculations.gas.carbon_monoxide import calculate_MCO_sec, calculate_MCO_year
from calculations.gas.methane import calculate_ethanethiol, calculate_methane

from gross_emissions.calculate_gross_emissions import calculate_gross_emissions

from prilozhenie.prilozhenie2.val_moving import generate_moving_desription
from prilozhenie.prilozhenie2.val_station import generate_station_desription

# Всем организованным источникам выбросов присваивают номера от 0001 до 5999, всем неорганизованным источникам - с 6001.
c = 100
def fill_document_into_table(num_of_boilers, fuel_type, num_of_days_in_work, num_of_hours_in_day_in_work, fuel_consumption_year, throughpout, tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk):
    
    generate_moving_desription(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    generate_station_desription(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    path = os.path.join(os.getcwd(), "calculations", "tables", "razdel2_table.docx")
    print(path)
    # doc = Document(r"calculations\tables\razdel2_table.docx")
    doc = Document(path)
    
    COgik_warmup, COgik_warm, CO_mik_warm, CO_mik_warmup, Gno2_warm, Gno2_warm_up, Gno_warm, Gno_warm_up, Mno2_warm, Mno_warm, Mno_warm_up, Mno2_warm_up, SOgik, SO_gik_warm, SO_mik, SO_mik_warm, petrol_gik_warm, petrol_gik_warmup, petrol_mik_warm, petrol_mik_warmup = calculate_gross_emissions(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    fuel_consumption_sec = int(fuel_consumption_year) * 1000 / (60 * 60 * 24 * num_of_days_in_work)
    table = doc.tables[0]

    new_row = table.add_row()
    row_idx = 1
    table.cell(row_idx, 0).text = "0307"
    table.cell(row_idx, 1).text = "Азота диоксид"
    table.cell(row_idx, 2).text = "0,2"
    table.cell(row_idx, 3).text = "0,1"
    table.cell(row_idx, 4).text = "0,04"
    table.cell(row_idx, 6).text = "7"
    azota_dioksid_sec = (calculate_MNO2_g(fuel_consumption_sec=fuel_consumption_sec) + Gno2_warm + Gno2_warm_up)
    azota_dioksid_year = calculate_Mno2_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec) + Mno_warm + Mno_warm_up
    table.cell(row_idx, 7).text = str(azota_dioksid_sec)
    table.cell(row_idx, 8).text = str(azota_dioksid_year)

    
    new_row = table.add_row()
    row_idx = 2
    table.cell(row_idx, 0).text = "0304"
    table.cell(row_idx, 1).text = "Азот (II) оксид"
    table.cell(row_idx, 2).text = "0,4"
    table.cell(row_idx, 4).text = "0,06"
    table.cell(row_idx, 6).text = "3"
    azot_2_oksid_sec = calculate_MNO_g(fuel_consumption_sec=fuel_consumption_sec) + Gno_warm + Gno_warm_up
    azot_2_oksid_year = calculate_MNO_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec) + Mno2_warm + Mno2_warm_up
    table.cell(row_idx, 7).text = str(azot_2_oksid_sec)
    table.cell(row_idx, 8).text = str(azot_2_oksid_year)
    
    new_row = table.add_row()
    row_idx = 3
    table.cell(row_idx, 0).text = "0330"
    table.cell(row_idx, 1).text = "Сера диоксид"
    table.cell(row_idx, 2).text = "0.5"
    table.cell(row_idx, 3).text = "0.05"
    table.cell(row_idx, 6).text = "4"
    sera_dioksid_sec = SOgik + SO_gik_warm
    sera_dioksid_year = SO_mik + SO_mik_warm
    table.cell(row_idx, 7).text = str(sera_dioksid_sec)
    table.cell(row_idx, 8).text = str(sera_dioksid_year)
    
    new_row = table.add_row()
    row_idx = 4 # Номер строки (начинается с 0)
    table.cell(row_idx, 0).text = "0337"
    table.cell(row_idx, 1).text = "Углерода оксид"
    table.cell(row_idx, 2).text = "5"
    table.cell(row_idx, 3).text = "3"
    table.cell(row_idx, 4).text = "3"
    table.cell(row_idx, 6).text = "4"
    uglerod_oksid_sec = calculate_MCO_sec(fuel_consumption_sec=fuel_consumption_sec) + COgik_warmup + COgik_warm
    uglerod_oksid_year = calculate_MCO_year(fuel_consumption_year=fuel_consumption_year) + CO_mik_warm + CO_mik_warmup
    table.cell(row_idx, 7).text = str(uglerod_oksid_sec)
    table.cell(row_idx, 8).text = str(uglerod_oksid_year)
    
    
    new_row = table.add_row()
    row_idx = 5
    table.cell(row_idx, 0).text = "0410"
    table.cell(row_idx, 1).text = "Метан"
    table.cell(row_idx, 5).text = "50"
    methane_sec, methane_year = calculate_methane(throughpout=throughpout)
    methane_sec_rounded = round(methane_sec, 6)
    table.cell(row_idx, 7).text = str(methane_sec_rounded)
    table.cell(row_idx, 8).text = str(methane_year)
    
    row_idx = 6
    col_idx = 0
    table.cell(row_idx, col_idx).text = "0703"
    table.cell(row_idx, 1).text = "бенз/а/пирен"
    table.cell(row_idx, 3).text = "0,000001"
    table.cell(row_idx, 4).text = "0,000001"
    table.cell(row_idx, 6).text = "1"
    benzapiren_sec = calculate_amount_of_benzapiren_sec(fuel_consumption_sec=fuel_consumption_sec)
    benzapiren_year = calculate_amount_of_benzapiren_year(fuel_consumption_year=fuel_consumption_year, num_of_days_in_work=num_of_days_in_work)
    table.cell(row_idx, 7).text = str(benzapiren_sec)
    table.cell(row_idx, 8).text = str(benzapiren_year)
     
    new_row = table.add_row()
    row_idx = 7
    table.cell(row_idx, 0).text = "1728"
    table.cell(row_idx, 1).text = "Этантиол"
    table.cell(row_idx, 6).text = "3"
    etantiol_sec, etantiol_year = calculate_ethanethiol(throughpout=throughpout)
    table.cell(row_idx, 7).text = str(etantiol_sec)
    table.cell(row_idx, 8).text = str(etantiol_year)

    new_row = table.add_row()
    row_idx = 8
    table.cell(row_idx, 0).text = "2704"
    table.cell(row_idx, 1).text = "Бензин (нефтяной, малосернистый)/ в пересчете на углерод"
    table.cell(row_idx, 2).text = "5"
    table.cell(row_idx, 3).text = "1.5"
    table.cell(row_idx, 6).text = "4"
    benzin_sec = petrol_gik_warmup + petrol_gik_warm
    benzin_year = petrol_mik_warm + petrol_mik_warmup
    table.cell(row_idx, 7).text = str(benzin_sec)
    table.cell(row_idx, 8).text = str(benzin_year)
    path_to_save = os.path.join(os.getcwd(), "calculations", "tables", "razdel2_table_filled.docx")
    doc.save(path_to_save)
    

    path = os.path.join(os.getcwd(), "calculations", "tables", "3_7.docx")
    print(path)
    # doc = Document("calculations/tables/3_7.docx")
    doc = Document(path)
    table = doc.tables[0]
    table.cell(3,0).text = "0301\n 0304\n 0330\n 0337\n 0410\n 0703\n 1728\n 2704"
    table.cell(3,1).text = "Азота диоксид\n Азот (II) оксид\n Сера диоксид\n Углерода оксид\n Метан\n Бензапирен\n Бензин (нефтяной, малосернистый) /в пересчете на углерод/"
    table.cell(3,2).text = f"{azota_dioksid_year}\n {azot_2_oksid_year}\n {sera_dioksid_year}\n {uglerod_oksid_year}\n {methane_year}\n {benzapiren_year}\n {etantiol_year}\n {benzin_year}"
    table.cell(3,3).text = f"{azota_dioksid_year}\n {azot_2_oksid_year}\n {sera_dioksid_year}\n {uglerod_oksid_year}\n {methane_year}\n {benzapiren_year}\n {etantiol_year}\n {benzin_year}"
    table.cell(3,4).text = f"{calculate_Mno2_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec)}\n {calculate_MNO_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec)}\n \n {calculate_MCO_year(fuel_consumption_year=fuel_consumption_year)}\n {methane_year}\n {benzapiren_year}\n {etantiol_year}\n "
    table.cell(3,9).text = f"{azota_dioksid_year}\n {azot_2_oksid_year}\n {sera_dioksid_year}\n {uglerod_oksid_year}\n {methane_year}\n {benzapiren_year}\n {etantiol_year}\n {benzin_year}"
    table.add_row()
    table.cell(4,0).merge(table.cell(4,1))
    table.cell(4,0).text = """ВСЕГО
                              в том числе"""
    
    table.cell(4,2).text = str(azota_dioksid_year + azot_2_oksid_year + sera_dioksid_year + uglerod_oksid_year + methane_year + benzapiren_year + etantiol_year + benzin_year)
    table.cell(4,3).text = str(azota_dioksid_year + azot_2_oksid_year + sera_dioksid_year + uglerod_oksid_year + methane_year + benzapiren_year + etantiol_year + benzin_year)
    table.cell(4,4).text = str(calculate_Mno2_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec) + calculate_MNO_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec) + calculate_MCO_year(fuel_consumption_year=fuel_consumption_year) +  methane_year + benzapiren_year + etantiol_year)
    table.cell(4,9).text = str(azota_dioksid_year + azot_2_oksid_year + sera_dioksid_year + uglerod_oksid_year + methane_year + benzapiren_year + etantiol_year + benzin_year)
    
    table.add_row()
    table.cell(5,0).merge(table.cell(5,1))
    table.cell(5,0).text = """ТВЕРДЫХ
                              """
    
    table.cell(5,2).text = str(benzapiren_year)
    table.cell(5,3).text = str(benzapiren_year)
    table.cell(5,4).text = str(benzapiren_year)
    table.cell(5,9).text = str(benzapiren_year)
    
    table.add_row()
    table.cell(6,0).merge(table.cell(6,1))
    table.cell(6,0).text = """Газообразных и жидких:
                              """
    
    table.cell(6,2).text = str(float( table.cell(4,2).text) - benzapiren_year)
    table.cell(6,3).text = str(float( table.cell(4,3).text) - benzapiren_year)
    table.cell(6,4).text = str(float( table.cell(4,4).text) - benzapiren_year)
    table.cell(6,9).text = str(float( table.cell(4,2).text) - benzapiren_year)
    
    path_to_save = os.path.join(os.getcwd(), "calculations", "tables", "razdel3", "3_7.docx")
    # doc.save("calculations/tables/razdel3/3_7.docx")
    doc.save(path_to_save)


def generate_docx(num_of_boilers, fuel_type, num_of_days_in_work, num_of_hours_in_day_in_work, fuel_consumption_year, throughpout, tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk, car, speed, num_of_days_in_work_car, num_of_hours_in_day_in_work_car, num_of_days_in_work_parking, num_of_hours_in_day_in_work_parking, num_of_parkings, num_of_days_in_work_candle, num_of_hours_in_day_in_work_candle, num_of_candles, car_amount):
    fill_document_into_table(num_of_boilers, fuel_type, num_of_days_in_work, num_of_hours_in_day_in_work, fuel_consumption_year, throughpout, tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    fuel_consumption_sec = int(fuel_consumption_year) * 1000 / (60 * 60 * 24 * num_of_days_in_work)

    path = os.path.join(os.getcwd(), "calculations", "tables", "razdel2_table_filled.docx")
    doc = Document(path)
    table = doc.tables[0]
    column_index = -1
    total_sum = 0
    for row in table.rows:
        cell = row.cells[column_index]
        try:
            value = float(cell.text.strip())
            total_sum += value
        except ValueError:
            print(f"Skipping non-numeric value: {cell.text}")
    table.add_row()
    table.cell(len(table.rows) - 1, 2).text = "В С Е Г О"
    table.cell(len(table.rows) - 1, len(table.rows[0].cells) - 1).text = str(total_sum)
    parent = table._element.getparent()
    tmp_zhid = 0.04698605186
    tmp_tverd = 0.0000001
    par = doc.add_paragraph(f"""Всего на производственной площадке выделяется {len(doc.tables[0].rows) - 1} загрязняющих веществ. Валовый выброс загрязняющих веществ составляет {total_sum} тонн в год (твёрдых – {tmp_tverd}; газообразных и жидких – {tmp_zhid}""")
    parent.insert(parent.index(table._element), par._element)
    par = doc.add_paragraph("Перечень нормативно-методических документов, использованных для расчета:")
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = par.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True

    LIST_OF_METHODS = ["Приказ Министерства природных ресурсов и экологии РФ №871 от 19 ноября 2021 года “Об утверждении Порядка проведения инвентаризации стационарных источников и выбросов вредных (загрязняющих) веществ в атмосферный воздух, корректировки ее данных, документирования и хранения данных, полученных в результате проведения таких инвентаризации и корректировки (с изменениями на 17 сентября 2019 года).",
                       """СанПиН 2.1.3684-21 "Санитарно-эпидемиологические требования к содержанию территорий городских и сельских поселений, к водным объектам, питьевой воде и питьевому водоснабжению, атмосферному воздуху, почвам, жилым помещениям, эксплуатации производственных, общественных помещений, организации и проведению санитарно-противоэпидемических (профилактических) мероприятий.""",
                       """СанПиН 1.2.3685-21 "Гигиенические нормативы и требования к обеспечению безопасности и (или) безвредности для человека факторов среды обитания".""",
                       """Методике проведения инвентаризации выбросов загрязняющих веществ в атмосферу для автотранспортных предприятий""",
                       """Методика проведения инвентаризации выбросов загрязняющих веществ в атмосферу для баз дорожной техники (расчетным методом). Москва, 1998 (с Дополнениями к методике проведения инвентаризации выбросов загрязняющих веществ в атмосферу для баз дорожной техники (расчетным методом Москва, 1999)"""
                       """«Методика определения выбросов загрязняющих веществ в атмосферу при сжигании топлива в котлах производительностью менее 30 т пара в час или менее 20 Гкал/час», Москва, 1999 г., с учетом методического письма НИИ Атмосфера N 335/33-07 от 17 мая 2000 г и изменений к ним (письмо НИИ Атмосферы N 838/33-07 от 11.09.2001)."""
]
    for method in LIST_OF_METHODS:
        paragraph = doc.add_paragraph(method, style='List Number')

    path_to_save = os.path.join(os.getcwd(), "inventarization_description", "izav_info", "part2_without_intro.docx")
    doc.save(path_to_save)

    COgik_warmup, COgik_warm, CO_mik_warm, CO_mik_warmup, Gno2_warm, Gno2_warm_up, Gno_warm, Gno_warm_up, Mno2_warm, Mno_warm, Mno_warm_up, Mno2_warm_up, SOgik, SO_gik_warm, SO_mik, SO_mik_warm, petrol_gik_warm, petrol_gik_warmup, petrol_mik_warm, petrol_mik_warmup = calculate_gross_emissions(tnp, txx1, txx2, L2_warmup, L1_warmup, ab, Nk, Dp, Tr, L1_warm, L2_warm, Nkv, Nkk)
    mno2_t = calculate_Mno2_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec)
    MNO2_g = calculate_MNO2_g(fuel_consumption_sec=fuel_consumption_sec)
    mco_sec = calculate_MCO_sec(fuel_consumption_sec=fuel_consumption_sec)
    mco_year = calculate_MCO_year(fuel_consumption_year=fuel_consumption_year)
    mno_g = calculate_MNO_g(fuel_consumption_sec=fuel_consumption_sec)
    azot_2_oksid_year = calculate_MNO_t(fuel_consumption_year=fuel_consumption_year, fuel_consumption_sec=fuel_consumption_sec)
    methane_sec, methane_year = calculate_methane(throughpout=throughpout)
    methane_sec = round(methane_sec, 6)
    benzapiren_sec = calculate_amount_of_benzapiren_sec(fuel_consumption_sec=fuel_consumption_sec)
    benzapiren_year = calculate_amount_of_benzapiren_year(fuel_consumption_year=fuel_consumption_year, num_of_days_in_work=num_of_days_in_work)
    etantiol_sec, etantiol_year = calculate_ethanethiol(throughpout=throughpout)
    sera_dioksid_sec = SOgik + SO_gik_warm
    sera_dioksid_year = SO_mik + SO_mik_warm

    fill_table_3_1(sera_dioksid_sec, sera_dioksid_year, Mno_warm, SO_mik_warm, CO_mik_warm, petrol_mik_warm, Gno_warm, SO_gik_warm, COgik_warm, petrol_gik_warm, mco_sec, mco_year, MNO2_g, azot_2_oksid_year, mno2_t, mno_g, benzapiren_sec,benzapiren_year, methane_sec, methane_year,etantiol_sec, etantiol_year, Gno2_warm, Gno2_warm_up, Mno2_warm, Mno2_warm_up, SOgik, SO_mik, COgik_warmup, CO_mik_warmup, petrol_gik_warmup, petrol_mik_warmup, num_of_boilers, num_of_days_in_work,num_of_hours_in_day_in_work, num_of_days_in_work_parking, num_of_hours_in_day_in_work_parking, num_of_parkings, num_of_days_in_work_candle, num_of_hours_in_day_in_work_candle, num_of_candles, num_of_days_in_work_car, num_of_hours_in_day_in_work_car, car_amount)
    fill_3_8_table(1, car, Nk, speed, fuel_type, num_of_days_in_work_car * num_of_hours_in_day_in_work_car,  num_of_days_in_work_car * num_of_hours_in_day_in_work_car, Mno2_warm, Mno_warm, SO_mik_warm, CO_mik_warm, petrol_mik_warm, Gno2_warm, Gno_warm, SO_gik_warm, COgik_warm, petrol_gik_warm)