import os
import sys
sys.path.append(os.getcwd())
import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from gross_emissions.gross_emission_lists_for_different_source import car, boiler_pipe, blow_off_plug

st.title("Информация об организации и источниках выбросов")

st.subheader("Введите информацию об организации")
organization_info = {
    "Численность рабочих по штатному": st.text_input("Численность рабочих по штатному", "21 человек"),
    "Среднесписочная численность": st.text_input("Среднесписочная численность", "17 человек"),
    "Дата ввода объекта в эксплуатацию": st.text_input("Дата ввода объекта в эксплуатацию", "2001 г. (Котельная)"),
    "Адрес расположения площадки": st.text_input("Адрес расположения площадки", "461987, Оренбургская область, Первомайский район, с. Мирошкино, ул. Центральная, 27"),
    "Земельный участок": st.text_input("Земельный участок", "56:22:0501001:4 площадью 15 550 м²"),
    "Категория земель": st.text_input("Категория земель", "Земли населенных пунктов"),
    "Виды разрешённого использования": st.text_input("Виды разрешённого использования", "для размещения административного здания и производственных помещений"),
    "Арендаторы, осуществляющие негативное воздействие на окружающую среду": st.text_input("Арендаторы, осуществляющие негативное воздействие на окружающую среду", "отсутствуют"),
    "Водоснабжение": st.text_input("Водоснабжение", "централизованное"),
    "Водоотведение": st.text_input("Водоотведение", "выгребная яма"),
    "Теплоснабжение": st.text_input("Теплоснабжение", "от собственной котельной"),
    "Перспективы развития": st.text_input("Перспективы развития", "в ближайшие 7 лет не запланированы"),
    "На производственной площадка расположены следующие источники загрязнения атмосферы:" : st.text_input("На производственной площадка расположены следующие источники загрязнения атмосферы:", "")
}

# Инициализация состояния для источников
if "sources" not in st.session_state:
    st.session_state.sources = []  # Список для хранения источников
if "next_source_id" not in st.session_state:
    st.session_state.next_source_id = 1  # Уникальный ID для источников

# Динамическое добавление источников
st.subheader("Добавить источник")
if st.button("Добавить источник"):
    # Добавляем новый источник с уникальным ID
    st.session_state.sources.append({"id": st.session_state.next_source_id, "type": "Организованный", "data": {}})
    st.session_state.next_source_id += 1

# Отображение и ввод данных для каждого источника
for source in st.session_state.sources:
    st.subheader(f"Источник №{source['id']}")

    # Выбор типа источника
    source_type = st.selectbox(f"Тип источника для источника №{source['id']}", ["", "Труба котельной", "Продувочная свеча", "Открытая стоянка"], key=f"type_{source['id']}")

    if source_type == "Труба котельной":
        # Ввод данных для трубы котельной
        source["data"]["Номер источника"] = st.text_input(f"Номер источника для источника №{source['id']}", f"ИЗАВ №{source['id']:04d} – Труба котельной")
        source["data"]["ИВ (001)"] = st.text_input("Название источника для заполнения ИВ", f"Котлы водогрейные КЧМ -5", key="my_unique_key1")
        source["data"]["Тип котла"] = st.text_input(f"Тип котла для источника №{source['id']}", "водогрейный КЧМ -5")
        source["data"]["Мощность котла"] = st.text_input(f"Мощность котла для источника №{source['id']}", "21 кВт")
        source["data"]["Количество котлов в работе"] = st.text_input(f"Количество котлов в работе для источника №{source['id']}", "2 в работе, 0 в резерве")
        source["data"]["Количество одновременно работающих котлов"] = st.text_input(f"Количество одновременно работающих котлов для источника №{source['id']}", "2 шт.")
        source["data"]["Вид топлива"] = st.text_input(f"Вид топлива для источника №{source['id']}", "газ горючий природный")
        source["data"]["Фактический расход топлива на один котёл"] = st.text_input(f"Фактический расход топлива на один котёл для источника №{source['id']}", "15 024 м³")
        source["data"]["Плотность газа"] = st.text_input(f"Плотность газа для источника №{source['id']}", "0,733 кг/м³")
        source["data"]["Резервное топливо"] = st.text_input(f"Резервное топливо для источника №{source['id']}", "отсутствует")
        source["data"]["Количество дней работы в год"] = st.text_input(f"Количество дней работы в год для источника №{source['id']}", "181 день")
        source["data"]["Время работы в сутки"] = st.text_input(f"Время работы в сутки для источника №{source['id']}", "24 часа/день")
        source["data"]["Высота дымовой трубы"] = st.text_input(f"Высота дымовой трубы для источника №{source['id']}", "6 м")
        source["data"]["Диаметр дымовой трубы"] = st.text_input(f"Диаметр дымовой трубы для источника №{source['id']}", "0,470 м")
        source["data"]["Загрязняющие вещества"] = st.text_input(f"Загрязняющие вещества для источника №{source['id']}", f"{boiler_pipe}")

    elif source_type == "Продувочная свеча":
        source["data"]["Номер источника"] = st.text_input(f"Номер источника для источника №{source['id']}", f"ИЗАВ №{source['id']:04d} – Продувочная свеча")
        source["data"]["ИВ (001)"] = st.text_input("Название источника", f"Продувочная свеча", key="my_unique_key2")
        source["data"]["Время работы"] = st.text_input(f"Время работы для источника №{source['id']}", "1,5 час/день, 1 дней/год")
        source["data"]["Длина продуваемого трубопровода"] = st.text_input(f"Длина продуваемого трубопровода для источника №{source['id']}", "5 м")
        source["data"]["Внутренний диаметр продуваемого трубопровода"] = st.text_input(f"Внутренний диаметр продуваемого трубопровода для источника №{source['id']}", "57 мм")
        source["data"]["Диаметр трубы продувочной свечи"] = st.text_input(f"Диаметр трубы продувочной свечи для источника №{source['id']}", "32 мм")
        source["data"]["Высота трубы продувочной свечи"] = st.text_input(f"Высота трубы продувочной свечи для источника №{source['id']}", "6 м")
        source["data"]["Загрязняющие вещества"] = st.text_input(f"Загрязняющие вещества для источника №{source['id']}", f"{blow_off_plug}")

    elif source_type == "Открытая стоянка":
        source["data"]["Номер источника"] = st.text_input(f"Номер источника для источника №{source['id']}", f"ИЗАВ №{source['id']:04d} – Открытая стоянка")
        source["data"]["ИВ (001)"] = st.text_input("Название", f"Открытая стоянка", key="my_unique_key3")
        source["data"]["Размеры стоянки"] = st.text_input(f"Размеры стоянки для источника №{source['id']}", "3 х 8 м")
        source["data"]["Марка оборудования"] = st.text_input(f"Марка оборудования для источника №{source['id']}", "ГАЗ 322171")
        source["data"]["Количество ТС"] = st.text_input(f"Количество ТС для источника №{source['id']}", "1")
        source["data"]["Тип ТС"] = st.text_input(f"Тип ТС для источника №{source['id']}", "автобус")
        source["data"]["Время работы"] = st.text_input(f"Время работы для источника №{source['id']}", "1,5 часа/день, 181 дней/год")
        source["data"]["Вид топлива"] = st.text_input(f"Вид топлива для источника №{source['id']}", "Бензин")
        source["data"]["Загрязняющие вещества"] = st.text_input(f"Загрязняющие вещества для источника №{source['id']}", f"{car}")

        st.subheader("Таблица для Открытой стоянки")
        parking_data = {
            "Марка оборудования": [source["data"]["Марка оборудования"]],
            "Количество ТС": [source["data"]["Количество ТС"]],
            "Тип ТС": [source["data"]["Тип ТС"]],
            "Время работы": [source["data"]["Время работы"]],
            "Вид топлива": [source["data"]["Вид топлива"]]
        }
        st.table(parking_data)

# Отображение введённой информации
st.subheader("Введённая информация об организации")
for key, value in organization_info.items():
    st.write(f"**{key}:** {value}")

st.subheader("Введённая информация об источниках")
for source in st.session_state.sources:
    for key, value in source["data"].items():
        st.write(f"**{key}:** {value}")

def generate_docx(organization_info, sources):
    doc = Document()
    # style = doc.styles['Normal']
    # style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading = doc.add_heading("1.2 Сведения о производственной деятельности, количестве, характеристиках и эффективности ГОУ", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    for key, value in organization_info.items():
        # doc.add_paragraph(f"{key}: {value}")
        p = doc.add_paragraph()
        p.add_run(f"{key}: {value}").bold = False   

    # Добавление информации об источниках
    for source in sources:
        heading = doc.add_heading(f"ИЗАВ №{source['id']} ({source['data'].get('Номер источника', '')})", level=3)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        print(source)
        # Если источник - "Открытая стоянка", добавляем текст перед таблицей
        if "Открытая стоянка" == source['data'].get('ИВ (001)', ''):
            par = doc.add_paragraph(f"ИВ (001): {source['data'].get('ИВ (001)', '')}")
            par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = par.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph(f"Размеры стоянки: {source['data'].get('Размеры стоянки', '')}")

            # Создаем таблицу для открытой стоянки
            table = doc.add_table(rows=2, cols=5)  # 2 строки: заголовок и данные
            table.style = 'Table Grid'

            # Заголовки таблицы
            headers = ["Марка оборудования", "Кол-во", "Тип ТС", "Время работы", "Вид топлива"]
            for i, header in enumerate(headers):
                table.cell(0, i).text = header

            # Данные для таблицы
            table.cell(1, 0).text = source["data"].get("Марка оборудования", "")
            table.cell(1, 1).text = source["data"].get("Количество ТС", "")
            table.cell(1, 2).text = source["data"].get("Тип ТС", "")
            table.cell(1, 3).text = source["data"].get("Время работы", "")
            table.cell(1, 4).text = source["data"].get("Вид топлива", "")

        # Добавляем остальные данные источника
        for key, value in source["data"].items():
            if key == "ИВ (001)":
                par = doc.add_paragraph(f"{key}: {value}")
                par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif key not in ["Марка оборудования", "Количество ТС", "Тип ТС", "Время работы", "Вид топлива", "Размеры стоянки", "Номер источника"]:
                doc.add_paragraph(f"{key}: {value}")

    file_path = "inventarization_description\izav_info\organization_sources_info.docx"
    doc.save(file_path)
    return file_path
# Кнопка для сохранения в DOCX
if st.button("Сохранить в DOCX"):
    file_path = generate_docx(organization_info, st.session_state.sources)
    with open(file_path, "rb") as f:
        st.download_button("Скачать DOCX", f, file_name="organization_sources_info.docx")