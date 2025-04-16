import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
doc = Document()
st.title("Характеристика пылегазоочистного оборудования и оценка его эффективности")

pgoy = st.text_input("Введите ПГОУ, установленное на вашем предприятии. Если такого нет, оставьте пустую строку")

if st.button("Отправить"):
    if pgoy:
        pass
    else:
        # 1.2.3 Характеристика пылегазоочистного оборудования и оценка его эффективности
        par = doc.add_paragraph()
        par.add_run('''1.2.3 Характеристика пылегазоочистного оборудования и оценка его эффективности''').bold = True
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("Пылегазоочистное оборудование (ПГОУ) на объекте негативного воздействия отсутствуют.")
        doc.save("object_property/dust_and_gas/dust_and_gus_info.docx")
        print("saved succesfully")
        st.success("Документ сохранён")