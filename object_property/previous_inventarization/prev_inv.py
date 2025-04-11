import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()
st.title("Сведения о проведённых инвентаризациях")

history = st.text_input("Введите информацию о проведённх инвентаризациях (если их не было оставьте поле пустым)")
org_name = st.text_input("Введите название вашей организации (сокращённое)", value="МБОУ «Мирошкинская СОШ»")

if st.button("Отправить"):
    if history:
        pass
    else:
        par = doc.add_paragraph()
        par.add_run("1.3 Сведения о результатах предыдущей инвентаризации").bold = True
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Ранее инвентаризации стационарных источников и выбросов вредных веществ в атмосферный воздух для  {org_name} не проводилась. ")
        doc.save("object_property/previous_inventarization/prev_inv_info.docx")
        st.success("Документ сохранён")