from docx import Document
from docx.shared import Pt

# Создаём документ
doc = Document()

# Добавляем таблицу (3 строки, 5 колонок)
table = doc.add_table(rows=3, cols=6)
table.style = 'Table Grid'

# Заполняем заголовок
table.cell(0, 0).merge(table.cell(0, 2))  # Объединение 1-3 колонок
table.cell(0, 3).merge(table.cell(0, 4))  # Объединение 4-5 колонок

table.cell(0, 0).text = "СанПиН 2.2.1/2.1.1.1200-03"
table.cell(0, 3).text = "Характер производства"
table.cell(0, 5).text = "Нормативный размер СЗЗ"

# Вторая строка заголовка
table.cell(1, 0).text = "Раздел*"
table.cell(1, 1).text = "класс опасности"
table.cell(1, 2).text = "пункт"
table.cell(1, 3).merge(table.cell(1, 4))  # Объединение 4-5 колонок

# Третья строка с содержимым
table.cell(2, 0).merge(table.cell(2, 5))  # Объединение 1-3 колонок
table.cell(2, 0).text = ("Для собственных котельных тепловой мощностью менее 200 Гкал, "
                         "работающих на твердом, жидком и газообразном топливе, "
                         "размер санитарно-защитной зоны не устанавливается.")
# table.cell(2, 0).merge(table.cell(2, 5))  # Объединение 1-3 колонок

# Настройка размера шрифта
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Сохраняем документ
doc.save("sanitary_zone/sanpin_table.docx")
