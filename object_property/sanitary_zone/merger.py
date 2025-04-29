from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_break(doc):
    """Inserts a page break into the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def combine_all_docx(filename_master,files_list):
    number_of_sections=len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        print("sth")
        doc_temp = Document_compose(files_list[i])
        composer.append(doc_temp)
    composer.save("object_property\sanitary_zone\sanitary_zone.docx")
files_to_merge = [
    "object_property\sanitary_zone\description.docx",
    "object_property\dust_and_gas\dust_and_gus_info.docx",
    "object_property\previous_inventarization\prev_inv_info.docx"
]
# filename_master="report.docx"
document = Document()
document.save("emptyfile.docx")
filename_master="emptyfile.docx"
combine_all_docx(filename_master,files_to_merge)